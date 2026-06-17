from pathlib import Path

from core.checker import check_completeness
from core.extractor import extract_from_upload
from core.filler import fill_templates
from core.pipeline import run_pipeline
from core.scanner import build_classified_catalog, scan_directory, summarize_classification
from core.utils import load_json


def test_pipeline_normal():
    root = Path(__file__).resolve().parents[1]
    upload = root / "data" / "upload" / "normal"
    output = root / "data" / "output" / "test_run"
    if not upload.exists():
        upload = root / "data" / "upload" / "custom" / "20260614_174849" / "临时用完删"
    result = run_pipeline(upload, output, use_llm=False)
    assert len(result.file_list) >= 1
    assert result.catalog_df is not None
    assert Path(result.report_path).exists()
    assert len(result.completeness_issues) >= 1
    assert "核心数据源" in result.catalog_by_category
    assert "申报原始模版" in result.catalog_by_category
    assert len(result.filled_files) >= 7


def test_file_classification():
    root = Path(__file__).resolve().parents[1]
    upload = root / "data" / "upload" / "custom" / "20260614_174849" / "临时用完删"
    if not upload.exists():
        return
    files = scan_directory(upload)
    counts = summarize_classification(files)
    assert counts["核心数据源"] >= 2
    assert counts["申报原始模版"] >= 5
    manual = [f for f in files if "说明书" in f.file_name]
    ch1 = [f for f in files if f.file_category == "申报原始模版"]
    assert manual and manual[0].file_category == "核心数据源"
    assert all("CH1" in f.file_name or "声明" in f.file_name or "申请表" in f.file_name for f in ch1[:3])


def test_completeness_cmde_rules():
    root = Path(__file__).resolve().parents[1]
    upload = root / "data" / "upload" / "custom" / "20260614_174849" / "临时用完删"
    if not upload.exists():
        return
    files = scan_directory(upload)
    issues = check_completeness(files)
    rules = load_json("nmpa_rules.json")
    assert rules.get("regulation", {}).get("detail_url")
    assert any(i.rule_id == "CMDE-A7" for i in issues)
    assert any(i.rule_id == "R-CH3-01" for i in issues)
    assert any(i.rule_id == "CH1.6" for i in issues)


def test_scheme_b_fill():
    root = Path(__file__).resolve().parents[1]
    upload = root / "data" / "upload" / "custom" / "20260614_174849" / "临时用完删"
    if not upload.exists():
        return
    out = root / "data" / "output" / "test_scheme_b"
    info = extract_from_upload(upload, use_llm=False)
    assert "新冠" in info.product_name or "2019-nCoV" in info.product_name
    filled = fill_templates(upload, out, info, file_list=scan_directory(upload))
    assert len(filled) >= 7
    from docx import Document

    ch14 = out / "CH1.4_申请表_已填写.docx"
    assert ch14.exists()
    doc = Document(str(ch14))
    t0 = doc.tables[0].rows
    assert info.product_name in t0[0].cells[1].text
    assert "6840-3-017" in t0[6].cells[1].text
    assert "☑第三类" in t0[1].cells[1].text
    assert "☑否" in t0[2].cells[1].text
    assert "☑临床试验" in doc.tables[1].rows[0].cells[1].text
    ch14_text = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "见附件" not in ch14_text
    assert "器审中心" not in ch14_text
    assert "符合率" not in doc.tables[1].rows[0].cells[1].text
    assert "临床数据库.zip" not in ch14_text
    assert info.manufacturer_name in ch14_text
    if info.manufacturer_address and info.manufacturer_address != "未提及":
        assert info.manufacturer_address[:8] in ch14_text

    ch15 = out / "CH1.5_产品列表_已填写.docx"
    assert ch15.exists()
    doc5 = Document(str(ch15))
    t0 = doc5.tables[0].rows[1]
    table_text = " ".join(c.text for r in doc5.tables[0].rows for c in r.cells)
    assert info.product_name.split("（")[0] in doc5.paragraphs[6].text or info.product_name in doc5.paragraphs[6].text
    assert "2019-nCoV" in t0.cells[3].text or "ORF1ab" in t0.cells[3].text
    assert "分管包装" in t0.cells[0].text
    assert "2019-nCoV阳性对照品" in table_text or "2019-nCoV 阳性对照品" in table_text
    assert "2019-nCoV" in table_text or "ORF1ab" in table_text
    assert "RSV" not in table_text and "MP、RSV" not in table_text
    assert "待补充" in t0.cells[1].text or t0.cells[1].text.isdigit() or t0.cells[1].text.startswith("601")
    t1 = doc5.tables[1].rows[1]
    assert "2019-nCoV" in t1.cells[0].text or "ORF1ab" in t1.cells[0].text or "nCoV" in t1.cells[0].text

    ch111 = out / "CH1.11.1_标准清单_已填写.docx"
    assert ch111.exists()
    t111 = "\n".join(p.text for p in Document(str(ch111)).paragraphs)
    assert info.product_name in t111
    assert t111.count("符合的标准清单如下") <= 1
    assert "系统自动匹配说明" in t111
    assert "指导原则" in t111 or "CMDE" in t111
    assert "2019-nCoV核酸检测试剂国家参考品" in t111 or "系统检索说明" in t111
    assert "2023年09月20日" not in t111

    ch115 = out / "CH1.11.5_真实性声明_已填写.docx"
    text = "\n".join(p.text for p in Document(str(ch115)).paragraphs)
    assert info.product_name in text
    assert "申请表" in text
    assert "产品说明书" in text
    assert "综述资料" not in text
    assert "主要原材料" not in text

    ch12 = out / "CH1.2_监管信息目录_已填写.docx"
    assert ch12.exists()
    d12 = Document(str(ch12))
    assert info.product_name in d12.paragraphs[1].text

    ch19 = out / "CH1.9_申报前沟通说明_已填写.doc"
    assert ch19.exists()
    t19 = ch19.read_bytes().decode("utf-16-le", errors="ignore")
    assert info.product_name in t19 or "2019-nCoV" in t19
