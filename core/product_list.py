from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement

from core.extractor import ExtractedInfo, read_docx_full_text
from core.llm_client import call_llm, should_use_llm
from core.utils import read_prompt


@dataclass
class ProductComponent:
    category: str
    name: str
    composition: str
    qty_24: str = ""
    qty_48: str = ""
    qty_96: str = ""


@dataclass
class ProductListData:
    pack_labels: dict[str, str] = field(default_factory=dict)
    catalog_numbers: dict[str, str] = field(default_factory=dict)
    spec_a: list[ProductComponent] = field(default_factory=list)
    spec_b: list[ProductComponent] = field(default_factory=list)
    comparison: list[dict[str, str]] = field(default_factory=list)
    source: str = "规则"

    def to_dict(self) -> dict[str, Any]:
        return {
            "pack_labels": self.pack_labels,
            "catalog_numbers": self.catalog_numbers,
            "spec_a": [c.__dict__ for c in self.spec_a],
            "spec_b": [c.__dict__ for c in self.spec_b],
            "comparison": self.comparison,
            "source": self.source,
        }


_PACK_KEYS = (
    "spec_a_24",
    "spec_a_48",
    "spec_a_96",
    "spec_b_24",
    "spec_b_48",
    "spec_b_96",
)

_CATEGORY_RULES = (
    (r"增强剂", "增强剂"),
    (r"阳性对照|阴性对照|对照品", "质控品"),
    (r"处理液|保存液", "处理液"),
    (r"反应液|PCR", "反应液"),
)


def _guess_category(name: str) -> str:
    for pat, cat in _CATEGORY_RULES:
        if re.search(pat, name, re.I):
            return cat
    return "其他"


def _norm_cell(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip())


def _is_covid_product(product_name: str) -> bool:
    return any(k in product_name for k in ("2019-nCoV", "新型冠状病毒", "新冠", "SARS-CoV-2"))


def _canonicalize_component_name(name: str, product_name: str = "") -> str:
    """统一组分命名：新冠类产品中 nCoV 阳性对照等须带 2019- 前缀。"""
    name = _norm_cell(name)
    if not name or not _is_covid_product(product_name):
        return name
    fixed = re.sub(r"(?<![\w-])nCoV(?=\s*阳性|\s*阴性|\s*PCR|阳性|阴性)", "2019-nCoV", name, flags=re.I)
    fixed = re.sub(r"2019-2019-", "2019-", fixed)
    return fixed


def _is_composition_table(header: list[str]) -> bool:
    joined = "".join(header)
    return "组分" in joined and "主要组成成分" in joined and "24人份" in joined


def _parse_composition_table(table: Any, *, product_name: str = "") -> list[ProductComponent]:
    rows = table.rows
    if len(rows) < 2:
        return []
    header = [_norm_cell(c.text) for c in rows[0].cells]
    if not _is_composition_table(header):
        return []

    qty_cols: dict[str, int] = {}
    for i, h in enumerate(header):
        if "24人份" in h:
            qty_cols["24"] = i
        elif "48人份" in h:
            qty_cols["48"] = i
        elif "96人份" in h:
            qty_cols["96"] = i

    name_col = next((i for i, h in enumerate(header) if "组分" in h), 0)
    comp_col = next((i for i, h in enumerate(header) if "主要组成成分" in h), 1)

    components: list[ProductComponent] = []
    for row in rows[1:]:
        cells = [_norm_cell(c.text) for c in row.cells]
        if not any(cells):
            continue
        name = _canonicalize_component_name(cells[name_col] if name_col < len(cells) else "", product_name)
        if not name or name in ("组分", "-", "/"):
            continue
        composition = cells[comp_col] if comp_col < len(cells) else ""
        components.append(
            ProductComponent(
                category=_guess_category(name),
                name=name,
                composition=composition,
                qty_24=cells[qty_cols["24"]] if "24" in qty_cols and qty_cols["24"] < len(cells) else "",
                qty_48=cells[qty_cols["48"]] if "48" in qty_cols and qty_cols["48"] < len(cells) else "",
                qty_96=cells[qty_cols["96"]] if "96" in qty_cols and qty_cols["96"] < len(cells) else "",
            )
        )
    return components


@dataclass
class PackWriteBlock:
    pack_label: str
    catalog: str
    size: int
    components: list[ProductComponent]


def _split_combined_pack_label(label: str) -> list[str]:
    """将「规格A：大包装：24人份/盒；分管包装：24人份/盒」拆成独立包装规格行。"""
    label = _norm_cell(label)
    if not label:
        return []
    if "；" not in label and ";" not in label:
        return [label]

    spec_m = re.match(r"(规格[AB])[：:]", label)
    spec_prefix = f"{spec_m.group(1)}：" if spec_m else ""
    out: list[str] = []
    for part in re.split(r"[；;]", label):
        part = part.strip()
        if not part:
            continue
        if part.startswith("规格"):
            out.append(part)
        elif spec_prefix and re.match(r"(大包装|分管包装)", part):
            out.append(spec_prefix + part)
        else:
            out.append(part)
    return out if out else [label]


def _catalog_for_pack(catalogs: dict[str, str], pack_key: str, sub_label: str) -> str:
    suffix = ""
    if "大包装" in sub_label and "分管包装" not in sub_label:
        suffix = "_bulk"
    elif "分管包装" in sub_label:
        suffix = "_tube"
    for key in (f"{pack_key}{suffix}", pack_key):
        val = catalogs.get(key)
        if val and str(val).lower() not in ("null", "none", "未提及", ""):
            return str(val).strip()
    return "待补充"


def _build_pack_write_blocks(data: ProductListData) -> list[PackWriteBlock]:
    blocks: list[PackWriteBlock] = []
    plan: list[tuple[str, int, list[ProductComponent]]] = [
        ("spec_a", 24, data.spec_a),
        ("spec_a", 48, data.spec_a),
        ("spec_a", 96, data.spec_a),
        ("spec_b", 24, data.spec_b or data.spec_a),
        ("spec_b", 48, data.spec_b or data.spec_a),
        ("spec_b", 96, data.spec_b or data.spec_a),
    ]
    for prefix, size, components in plan:
        if not components:
            continue
        pack_key = f"{prefix}_{size}"
        default_spec = "规格A" if prefix == "spec_a" else "规格B"
        raw_label = data.pack_labels.get(pack_key) or f"{default_spec}：{size}人份/盒"
        for sub_label in _split_combined_pack_label(raw_label):
            blocks.append(
                PackWriteBlock(
                    pack_label=sub_label,
                    catalog=_catalog_for_pack(data.catalog_numbers, pack_key, sub_label),
                    size=size,
                    components=components,
                )
            )
    return blocks


def _new_table_row(n_cols: int = 6) -> Any:
    tr = OxmlElement("w:tr")
    for _ in range(n_cols):
        tc = OxmlElement("w:tc")
        tc.append(OxmlElement("w:p"))
        tr.append(tc)
    return tr


def _rebuild_product_list_data_rows(table: Any, data_row_count: int) -> None:
    """清除模版合并单元格后按所需行数重建数据区，避免大/分管包装分行时错位。"""
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    n_cols = len(table.rows[0].cells) if table.rows else 6
    for _ in range(data_row_count):
        table._tbl.append(_new_table_row(n_cols))


def _merge_pack_cell_groups(table: Any, ranges: list[tuple[int, int]]) -> None:
    """每组包装规格合并「包装规格」「货号」列，与监管模版呈现一致。"""
    for start, end in ranges:
        if start >= end:
            continue
        for col in (0, 1):
            top = table.cell(start, col)
            for row_i in range(start + 1, end + 1):
                top.merge(table.cell(row_i, col))


def _parse_pack_labels(pack_specs: str) -> dict[str, str]:
    labels = {k: "" for k in _PACK_KEYS}
    defaults = {
        "spec_a_24": "规格A：24人份/盒",
        "spec_a_48": "规格A：48人份/盒",
        "spec_a_96": "规格A：96人份/盒",
        "spec_b_24": "规格B：24人份/盒",
        "spec_b_48": "规格B：48人份/盒",
        "spec_b_96": "规格B：96人份/盒",
    }
    if not pack_specs or pack_specs == "未提及":
        return defaults

    for spec, letter in (("规格A", "a"), ("规格B", "b")):
        m = re.search(rf"{spec}[：:](.*?)(?=规格[AB]|$)", pack_specs)
        if not m:
            continue
        block = m.group(1)
        bulk_m = re.search(r"大包装[：:]([^；;]+)", block)
        tube_m = re.search(r"分管包装[：:]([^；;]+)", block)
        for size in ("24", "48", "96"):
            size_pat = rf"{size}\s*人份/盒"
            if not re.search(size_pat, block):
                continue
            parts: list[str] = []
            if bulk_m and re.search(size_pat, bulk_m.group(1)):
                parts.append(f"大包装：{size}人份/盒")
            if tube_m and re.search(size_pat, tube_m.group(1)):
                parts.append(f"分管包装：{size}人份/盒")
            if parts:
                labels[f"spec_{letter}_{size}"] = f"{spec}：" + "；".join(parts)
            else:
                labels[f"spec_{letter}_{size}"] = f"{spec}：{size}人份/盒"

    for k, v in defaults.items():
        if not labels.get(k):
            labels[k] = v
    return labels


def _extract_catalog_from_text(text: str) -> dict[str, str]:
    found: dict[str, str] = {}
    for m in re.finditer(
        r"(规格A|规格B)\s*[：:\s]*(大包装|分管包装)?[：:\s]*(\d+)\s*人份/盒[^\d]{0,30}(601\d{7,8}|\d{10})",
        text,
    ):
        letter = "a" if "A" in m.group(1) else "b"
        size = m.group(3)
        pack_type = m.group(2) or ""
        key = f"spec_{letter}_{size}"
        if pack_type == "大包装":
            key += "_bulk"
        elif pack_type == "分管包装":
            key += "_tube"
        found[key] = m.group(4)
    for m in re.finditer(r"货号[：:\s]*(\d{10,12})", text):
        # single catalog without pack mapping — skip
        pass
    return found


def _norm_component_name(name: str) -> str:
    n = _norm_cell(name).lower().replace(" ", "")
    n = re.sub(r"2019[-]?ncov", "ncov", n)
    return n


def _build_comparison(spec_a: list[ProductComponent], spec_b: list[ProductComponent]) -> list[dict[str, str]]:
    a_map = {_norm_component_name(c.name): c for c in spec_a}
    b_map = {_norm_component_name(c.name): c for c in spec_b}
    order = list(dict.fromkeys([_norm_component_name(c.name) for c in spec_a] + [_norm_component_name(c.name) for c in spec_b]))
    rows: list[dict[str, str]] = []

    for key in order:
        a_comp = a_map.get(key)
        b_comp = b_map.get(key)
        if not a_comp and not b_comp:
            continue
        display_name = (a_comp or b_comp).name
        a_text = a_comp.composition if a_comp else "/"
        b_text = b_comp.composition if b_comp else "/"
        if a_comp and b_comp:
            same = "相同" if _norm_cell(a_comp.composition) == _norm_cell(b_comp.composition) else "不同"
        else:
            same = "不同"
        rows.append(
            {
                "name": display_name,
                "spec_a_composition": a_text,
                "spec_b_composition": b_text,
                "same_or_diff": same,
            }
        )
    return rows


def extract_product_list_by_rules(manual: Path, info: ExtractedInfo) -> ProductListData:
    doc = Document(str(manual))
    product_name = info.product_name if info.product_name != "未提及" else ""
    comp_tables: list[list[ProductComponent]] = []
    for table in doc.tables:
        parsed = _parse_composition_table(table, product_name=product_name)
        if parsed:
            comp_tables.append(parsed)

    data = ProductListData()
    data.pack_labels = _parse_pack_labels(info.pack_specs)
    data.catalog_numbers = _extract_catalog_from_text(read_docx_full_text(manual))

    if comp_tables:
        data.spec_a = comp_tables[0]
        data.spec_b = comp_tables[2] if len(comp_tables) >= 3 else (comp_tables[1] if len(comp_tables) >= 2 else comp_tables[0])
        if len(comp_tables) >= 3 and any("增强剂" in c.name for c in comp_tables[2]):
            data.spec_b = comp_tables[2]
        elif any("增强剂" in c.name for c in (comp_tables[-1] if comp_tables else [])):
            data.spec_b = comp_tables[-1]

    data.comparison = _build_comparison(data.spec_a, data.spec_b)
    data.source = "规则"
    return data


def _merge_llm_product_list(data: ProductListData, llm: dict[str, Any], *, product_name: str = "") -> ProductListData:
    if not llm:
        return data

    for key in _PACK_KEYS:
        val = (llm.get("pack_labels") or {}).get(key)
        if val:
            data.pack_labels[key] = str(val).strip()

    for key in _PACK_KEYS:
        val = (llm.get("catalog_numbers") or {}).get(key)
        if val and str(val).lower() not in ("null", "none", "未提及", ""):
            data.catalog_numbers[key] = str(val).strip()

    def _load_components(items: Any) -> list[ProductComponent]:
        if not isinstance(items, list):
            return []
        out: list[ProductComponent] = []
        for item in items:
            if not isinstance(item, dict):
                continue
            name = _canonicalize_component_name(str(item.get("name", "")).strip(), product_name)
            if not name:
                continue
            out.append(
                ProductComponent(
                    category=str(item.get("category") or _guess_category(name)).strip(),
                    name=name,
                    composition=str(item.get("composition", "")).strip(),
                    qty_24=str(item.get("qty_24", "")).strip(),
                    qty_48=str(item.get("qty_48", "")).strip(),
                    qty_96=str(item.get("qty_96", "")).strip(),
                )
            )
        return out

    llm_a = _load_components(llm.get("spec_a_components"))
    llm_b = _load_components(llm.get("spec_b_components"))
    if llm_a:
        data.spec_a = llm_a
    if llm_b:
        data.spec_b = llm_b

    notes = llm.get("comparison_notes")
    if isinstance(notes, list) and notes:
        data.comparison = [
            {
                "name": str(n.get("name", "")).strip(),
                "spec_a_composition": str(n.get("spec_a_composition", "")).strip(),
                "spec_b_composition": str(n.get("spec_b_composition", "")).strip(),
                "same_or_diff": str(n.get("same_or_diff", "不同")).strip(),
            }
            for n in notes
            if isinstance(n, dict) and n.get("name")
        ]
    elif data.spec_a or data.spec_b:
        data.comparison = _build_comparison(data.spec_a, data.spec_b)

    data.source = "规则+LLM" if (llm_a or llm_b) else data.source + "+LLM"
    return data


def extract_product_list_with_llm(text: str) -> dict[str, Any]:
    purpose = "阶段3 · 提取 CH1.5 产品列表表格数据"
    request_brief = (
        "从说明书【主要组成成分】等章节提取规格A/B各组分、主要组成成分、规格数量、货号；"
        "找不到的货号填 null，禁止编造。"
    )
    prompt = read_prompt("product_list_prompt.md", TEXT=text[:14000])
    raw = call_llm(
        prompt,
        purpose=purpose,
        request_brief=request_brief,
        prompt_file="prompts/product_list_prompt.md",
    )
    if not raw:
        return {}
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[-1].rsplit("```", 1)[0]
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {}


def extract_product_list(
    manual: Path,
    info: ExtractedInfo,
    *,
    use_llm: bool = True,
) -> ProductListData:
    data = extract_product_list_by_rules(manual, info)

    from core.run_logger import get_run_logger

    logger = get_run_logger()
    if logger:
        logger.python_only(
            "阶段3 · 规则解析产品列表",
            f"规格A {len(data.spec_a)} 个组分、规格B {len(data.spec_b)} 个组分；"
            f"货号 {len(data.catalog_numbers)} 条（来自说明书原文）。",
        )

    if should_use_llm(use_llm):
        text = read_docx_full_text(manual)
        llm_data = extract_product_list_with_llm(text)
        if llm_data:
            data = _merge_llm_product_list(data, llm_data, product_name=info.product_name)
            if logger:
                logger.python_only(
                    "阶段3 · 大模型补充产品列表",
                    f"合并后规格A {len(data.spec_a)} 项、规格B {len(data.spec_b)} 项，来源={data.source}。",
                )
        elif logger:
            logger.python_only("阶段3 · 产品列表大模型未生效", "沿用规则解析的组成表格。")

    return data


def _qty_for_size(comp: ProductComponent, size: int | str) -> str:
    key = str(size)
    return {"24": comp.qty_24, "48": comp.qty_48, "96": comp.qty_96}.get(key, "")


def _load_product_list_from_info(info: ExtractedInfo) -> ProductListData:
    raw = info.product_list or {}
    if not raw:
        return ProductListData()

    def _comps(items: Any) -> list[ProductComponent]:
        if not isinstance(items, list):
            return []
        out: list[ProductComponent] = []
        for item in items:
            if not isinstance(item, dict):
                continue
            out.append(
                ProductComponent(
                    category=str(item.get("category", "")),
                    name=str(item.get("name", "")),
                    composition=str(item.get("composition", "")),
                    qty_24=str(item.get("qty_24", "")),
                    qty_48=str(item.get("qty_48", "")),
                    qty_96=str(item.get("qty_96", "")),
                )
            )
        return out

    return ProductListData(
        pack_labels=dict(raw.get("pack_labels") or {}),
        catalog_numbers=dict(raw.get("catalog_numbers") or {}),
        spec_a=_comps(raw.get("spec_a")),
        spec_b=_comps(raw.get("spec_b")),
        comparison=list(raw.get("comparison") or []),
        source=str(raw.get("source") or "规则"),
    )


def apply_product_list_to_doc(
    doc: Document,
    info: ExtractedInfo,
    data: ProductListData | None = None,
) -> tuple[list[str], list[str]]:
    """将产品列表数据写入 CH1.5 两个表格。"""
    data = data or _load_product_list_from_info(info)
    written: list[str] = []
    skipped: list[str] = []

    if not doc.tables:
        return written, ["CH1.5 文档无表格"]

    product = info.product_name if info.product_name != "未提及" else ""
    if product:
        for para in doc.paragraphs:
            t = para.text.strip()
            if not t or "CH1.5" in t or "产品列表" in t or "监管信息" in t:
                continue
            if t in product or product in t:
                continue
            if ("呼吸道" in t or "肺炎支" in t) and len(t) < 100:
                if not (t.endswith("）") and "PCR" in t):
                    para.text = product.split("（")[0] if "（" in product else product
                    written.append("封面产品名称段落")
                    break

    if len(doc.tables) >= 1 and data.spec_a:
        table = doc.tables[0]
        write_blocks = _build_pack_write_blocks(data)
        data_row_count = sum(len(block.components) for block in write_blocks)
        _rebuild_product_list_data_rows(table, data_row_count)

        row_idx = 1
        filled_rows = 0
        block_ranges: list[tuple[int, int]] = []
        for block in write_blocks:
            block_start = row_idx
            for comp_idx, comp in enumerate(block.components):
                row = table.rows[row_idx]
                if len(row.cells) >= 6:
                    if comp_idx == 0:
                        row.cells[0].text = block.pack_label
                        row.cells[1].text = block.catalog
                    else:
                        row.cells[0].text = ""
                        row.cells[1].text = ""
                    row.cells[2].text = comp.category
                    row.cells[3].text = _canonicalize_component_name(comp.name, product)
                    row.cells[4].text = comp.composition
                    row.cells[5].text = _qty_for_size(comp, block.size)
                    filled_rows += 1
                row_idx += 1
            block_ranges.append((block_start, row_idx - 1))

        _merge_pack_cell_groups(table, block_ranges)

        pack_groups = len(write_blocks)
        written.append(
            f"主表 {filled_rows} 行（{pack_groups} 组包装规格 × 组分，大包装/分管包装分行）"
        )

    if len(doc.tables) >= 2 and data.comparison:
        table1 = doc.tables[1]
        for i, item in enumerate(data.comparison, start=1):
            if i >= len(table1.rows):
                break
            row = table1.rows[i]
            if len(row.cells) >= 4:
                row.cells[0].text = item.get("name", "")
                row.cells[1].text = item.get("spec_a_composition", "")
                row.cells[2].text = item.get("spec_b_composition", "")
                row.cells[3].text = item.get("same_or_diff", "")
        for i in range(len(data.comparison) + 1, len(table1.rows)):
            row = table1.rows[i]
            for cell in row.cells:
                cell.text = ""
        written.append(f"对比表 {len(data.comparison)} 行（规格A/B异同）")
    elif not data.comparison:
        skipped.append("对比表（未解析到规格A/B异同）")

    if not data.spec_a:
        skipped.append("主表（说明书未解析到【主要组成成分】表格）")

    return written, skipped
