from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pandas as pd

from core.utils import chapter_code_from_name, ensure_dir, load_json, normalize_text


@dataclass
class FileInfo:
    file_name: str
    file_path: str
    file_format: str
    page_count: int
    file_size_kb: float
    chapter_code: str | None
    title: str
    file_category: str = "未分类"


CATEGORY_LABELS = ("核心数据源", "申报原始模版", "输出文档模版")

SUPPORTED_SCAN_SUFFIXES = {".docx", ".pdf", ".doc"}


def count_docx_pages(path: Path) -> int:
    try:
        from docx import Document

        doc = Document(str(path))
        pages = sum(1 for p in doc.paragraphs if "PAGE" in p.text.upper() or "第" in p.text and "页" in p.text)
        if pages == 0:
            pages = max(1, len(doc.paragraphs) // 40 + len(doc.tables) // 2)
        return max(1, pages)
    except Exception:
        return 1


def count_pdf_pages(path: Path) -> int:
    try:
        import fitz

        with fitz.open(str(path)) as doc:
            return len(doc)
    except Exception:
        return 1


def count_pages(path: Path) -> int:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return count_pdf_pages(path)
    if suffix == ".docx":
        return count_docx_pages(path)
    return 1


def guess_title(name: str, chapter: str | None) -> str:
    mapping = {
        "CH1.2": "章节目录",
        "CH1.4": "申请表",
        "CH1.5": "产品列表",
        "CH1.9": "申报前沟通说明",
        "CH1.11.1": "符合标准清单",
        "CH1.11.5": "真实性声明",
        "CH1.11.6": "符合性声明",
    }
    if chapter and chapter in mapping:
        return mapping[chapter]
    if "说明书" in name:
        return "产品说明书"
    if "附件" in name or "申报资料要求" in name:
        return "法规参考文件"
    if "_已填写" in name:
        return name.replace("_已填写", "").replace(".docx", "")
    return Path(name).stem


def _load_catalog_rules() -> dict[str, Any]:
    return load_json("file_catalog.json")


def _match_patterns(name: str, patterns: list[str]) -> bool:
    norm = normalize_text(name)
    for pat in patterns:
        if normalize_text(pat) in norm:
            return True
    return False


def classify_upload_file(file_name: str, rules: dict[str, Any] | None = None) -> str:
    """规则引擎归类上传文件（流水线「Agent」的文件分类步骤）。"""
    rules = rules or _load_catalog_rules()
    for rule in rules.get("upload_rules", []):
        if _match_patterns(file_name, rule.get("patterns", [])):
            return rule["category"]
    return rules.get("default_upload_category", "核心数据源")


def classify_output_file(file_name: str, rules: dict[str, Any] | None = None) -> str:
    rules = rules or _load_catalog_rules()
    out_cfg = rules.get("output_rules", {})
    patterns = out_cfg.get("filename_patterns", [])
    if _match_patterns(file_name, patterns) or "_已填写" in file_name:
        return out_cfg.get("category", "输出文档模版")
    return out_cfg.get("category", "输出文档模版")


def scan_directory(upload_dir: Path, *, recursive: bool = True) -> list[FileInfo]:
    files: list[FileInfo] = []
    if not upload_dir.exists():
        return files

    catalog_rules = _load_catalog_rules()

    if recursive:
        paths = sorted(
            p for p in upload_dir.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_SCAN_SUFFIXES
        )
    else:
        paths = sorted(
            p for p in upload_dir.iterdir() if p.is_file() and p.suffix.lower() in SUPPORTED_SCAN_SUFFIXES
        )

    for path in paths:
        try:
            rel_name = str(path.relative_to(upload_dir))
        except ValueError:
            rel_name = path.name
        chapter = chapter_code_from_name(path.name)
        files.append(
            FileInfo(
                file_name=rel_name,
                file_path=str(path),
                file_format=path.suffix.lower().lstrip(".") or "unknown",
                page_count=count_pages(path),
                file_size_kb=round(path.stat().st_size / 1024, 1),
                chapter_code=chapter,
                title=guess_title(path.name, chapter),
                file_category=classify_upload_file(rel_name, catalog_rules),
            )
        )
    return files


def file_info_from_output_path(path: Path, *, base_label: str = "输出目录") -> FileInfo:
    name = path.name
    chapter = chapter_code_from_name(name)
    return FileInfo(
        file_name=f"{base_label}/{name}",
        file_path=str(path),
        file_format=path.suffix.lower().lstrip(".") or "unknown",
        page_count=count_pages(path) if path.exists() else 0,
        file_size_kb=round(path.stat().st_size / 1024, 1) if path.exists() else 0,
        chapter_code=chapter,
        title=guess_title(name, chapter),
        file_category="输出文档模版",
    )


def append_output_files(file_list: list[FileInfo], output_paths: list[str | Path]) -> list[FileInfo]:
    """将流水线生成的已填写文档并入清单（输出文档模版）。"""
    combined = list(file_list)
    seen = {f.file_path for f in combined}
    for raw in output_paths:
        path = Path(raw)
        if not path.exists() or str(path) in seen:
            continue
        combined.append(file_info_from_output_path(path))
        seen.add(str(path))
    return combined


def build_catalog_dataframe(file_list: list[FileInfo]) -> pd.DataFrame:
    rows = []
    for f in file_list:
        rows.append(
            {
                "文件分类": f.file_category,
                "RPS目录": f.chapter_code or "-",
                "标题": f.title,
                "适用情况": "R" if f.chapter_code else "O",
                "资料名称": f.file_name,
                "页码": f.page_count,
                "文件格式": f.file_format,
                "大小(KB)": f.file_size_kb,
            }
        )
    return pd.DataFrame(rows)


def build_classified_catalog(
    file_list: list[FileInfo],
) -> dict[str, pd.DataFrame]:
    """按三类分组返回目录表。"""
    groups: dict[str, list[FileInfo]] = {label: [] for label in CATEGORY_LABELS}
    for f in file_list:
        cat = f.file_category if f.file_category in groups else "核心数据源"
        groups[cat].append(f)

    result: dict[str, pd.DataFrame] = {}
    for label in CATEGORY_LABELS:
        if groups[label]:
            df = build_catalog_dataframe(groups[label]).drop(columns=["文件分类"])
            result[label] = df
        else:
            result[label] = pd.DataFrame(
                columns=["RPS目录", "标题", "适用情况", "资料名称", "页码", "文件格式", "大小(KB)"]
            )
    return result


def build_combined_catalog_dataframe(file_list: list[FileInfo]) -> pd.DataFrame:
    """带分类列的合并表（Excel 单 sheet 用）。"""
    return build_catalog_dataframe(file_list)


def export_ch12_excel(file_list: list[FileInfo], output_path: Path) -> Path:
    ensure_dir(output_path.parent)
    classified = build_classified_catalog(file_list)
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        build_combined_catalog_dataframe(file_list).to_excel(writer, index=False, sheet_name="全部文件")
        for label, df in classified.items():
            df.to_excel(writer, index=False, sheet_name=label[:31])
    return output_path


def files_index(file_list: list[FileInfo]) -> str:
    return "|".join(normalize_text(f.file_name) for f in file_list)


def summarize_classification(file_list: list[FileInfo]) -> dict[str, int]:
    counts = {label: 0 for label in CATEGORY_LABELS}
    for f in file_list:
        cat = f.file_category if f.file_category in counts else "核心数据源"
        counts[cat] += 1
    return counts
