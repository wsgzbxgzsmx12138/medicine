from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

from core.extractor import ExtractedInfo
from core.scanner import FileInfo
from core.utils import ensure_dir, glob_files, load_json, normalize_text


def _clip(s: str, n: int = 50) -> str:
    s = str(s).strip()
    return s if len(s) <= n else s[: n - 1] + "…"


def _log_doc_fill(
    output: Path,
    template: Path,
    fill_mode: str,
    *,
    written: list[str] | None = None,
    skipped: list[str] | None = None,
    note: str = "",
) -> None:
    from core.run_logger import get_run_logger

    logger = get_run_logger()
    if logger:
        logger.doc_fill(
            output.name,
            template_name=template.name,
            fill_mode=fill_mode,
            written=written,
            skipped=skipped,
            note=note,
        )


def _signing_date(cfg: dict[str, Any] | None = None) -> str:
    """声明/沟通类文档签署日期：默认使用本次运行当天。"""
    cfg = cfg or {}
    if cfg.get("signing_date"):
        return str(cfg["signing_date"])
    mapping = load_json("field_mapping.json")
    if mapping.get("signing_date"):
        return str(mapping["signing_date"])
    if mapping.get("use_run_date_for_signing", True):
        return datetime.now().strftime("%Y年%m月%d日")
    return datetime.now().strftime("%Y年%m月%d日")


def _template_signing_dates(cfg: dict[str, Any] | None = None) -> list[str]:
    cfg = cfg or {}
    if cfg.get("template_signing_dates"):
        return list(cfg["template_signing_dates"])
    return list(load_json("field_mapping.json").get("template_signing_dates", []))


def _replace_signing_dates_docx(doc: Document, cfg: dict[str, Any]) -> list[str]:
    """将模版占位签署日期替换为最新日期（不影响正文中的历史事件日期）。"""
    new_date = _signing_date(cfg)
    written: list[str] = []
    for old in _template_signing_dates(cfg):
        if not old or old == new_date:
            continue
        for para in doc.paragraphs:
            if old in para.text:
                para.text = para.text.replace(old, new_date)
                written.append(f"签署日期 {old} → {new_date}")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old in cell.text:
                        cell.text = cell.text.replace(old, new_date)
                        written.append(f"签署日期(表) {old} → {new_date}")
    return written


def _replace_signing_dates_text(text: str, cfg: dict[str, Any]) -> tuple[str, list[str]]:
    new_date = _signing_date(cfg)
    written: list[str] = []
    for old in _template_signing_dates(cfg):
        if not old or old == new_date or old not in text:
            continue
        count = text.count(old)
        text = text.replace(old, new_date)
        written.append(f"签署日期 {old} → {new_date}（×{count}）")
    return text, written


def _audit_stale_signing_dates(path: Path, cfg: dict[str, Any]) -> list[str]:
    """检查输出文件是否仍残留模版签署日期。"""
    stale = _template_signing_dates(cfg)
    if not stale or path.suffix.lower() not in (".docx", ".doc"):
        return []
    if path.suffix.lower() == ".docx":
        text = read_docx_plain_text(path)
    else:
        text = path.read_bytes().decode("utf-16-le", errors="ignore")
    return [d for d in stale if d in text]


def read_docx_plain_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _norm(s: str) -> str:
    return re.sub(r"\s+", "", s)


def _val(info: ExtractedInfo, key: str) -> str:
    v = getattr(info, key, None)
    if v is None or v == "未提及" or v == "" or v == []:
        return ""
    if isinstance(v, list):
        return "、".join(str(x) for x in v)
    return str(v).strip()


def _set_table_cell_by_label(
    doc: Document,
    row_label: str,
    value: str,
    *,
    table_index: int | None = None,
    section_label: str | None = None,
    max_len: int | None = None,
) -> bool:
    if not value:
        return False
    if max_len:
        value = value[:max_len]

    label_norm = _norm(row_label)
    section_norm = _norm(section_label) if section_label else None
    tables = doc.tables if table_index is None else [doc.tables[table_index]]

    for table in tables:
        in_section = section_norm is None
        for row in table.rows:
            cells = row.cells
            if not cells:
                continue
            row_text = _norm("".join(c.text for c in cells))
            if section_norm and section_norm in row_text and "名称" not in row_text:
                in_section = section_norm in row_text
            if section_norm and not in_section:
                continue
            if label_norm in row_text:
                for i, cell in enumerate(cells):
                    if label_norm in _norm(cell.text) and i + 1 < len(cells):
                        cells[i + 1].text = value
                        return True
                if len(cells) >= 2 and label_norm in _norm(cells[0].text):
                    cells[1].text = value
                    return True
                if len(cells) >= 3 and label_norm in _norm(cells[1].text):
                    cells[2].text = value
                    return True
    return False


def _replace_in_paragraphs(doc: Document, old: str, new: str) -> int:
    count = 0
    for para in doc.paragraphs:
        if old in para.text:
            para.text = para.text.replace(old, new)
            count += 1
    return count


def _replace_in_paragraph_range(
    doc: Document,
    old: str,
    new: str,
    *,
    start: int = 0,
    end: int | None = None,
) -> int:
    end = end if end is not None else len(doc.paragraphs)
    count = 0
    for para in doc.paragraphs[start:end]:
        if old in para.text:
            para.text = para.text.replace(old, new)
            count += 1
    return count


def _is_cover_product_fragment(text: str) -> bool:
    t = text.strip()
    if not t or t in ("申请人：", "申请人"):
        return False
    compact = re.sub(r"\s+", "", t)
    if len(t) <= 48 and any(k in compact for k in ("PCR", "法）", "原体", "合胞", "支原体", "RSV", "MP")):
        return True
    if "核酸检测试剂盒" in compact and len(t) <= 80:
        return True
    return False


def _replace_in_tables(doc: Document, old: str, new: str) -> int:
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old in cell.text:
                    cell.text = cell.text.replace(old, new)
                    count += 1
    return count


def _detection_short_name(info: ExtractedInfo) -> str:
    if info.targets:
        return "&".join(info.targets[:3])
    if "新冠" in info.product_name or "2019-nCoV" in info.product_name:
        return "2019-nCoV"
    name = info.product_name
    m = re.search(r"（(.+?)）", name)
    return m.group(1) if m else name[:20]


def _replace_mock_product(doc: Document, info: ExtractedInfo, cfg: dict[str, Any]) -> None:
    mapping = load_json("field_mapping.json")
    patterns = cfg.get("mock_patterns") or mapping.get("mock_product_patterns", [])
    short_patterns = mapping.get("mock_short_patterns", [])
    product = _val(info, "product_name")
    if not product:
        return

    for pat in patterns:
        _replace_in_paragraphs(doc, pat, product)
        _replace_in_tables(doc, pat, product)

    short = _detection_short_name(info)
    for pat in short_patterns:
        _replace_in_paragraphs(doc, pat, short)
        _replace_in_tables(doc, pat, short)


def _find_file_for_chapter(file_list: list[FileInfo], chapter_code: str) -> FileInfo | None:
    code_norm = normalize_text(chapter_code)
    for fi in file_list:
        if fi.chapter_code and normalize_text(fi.chapter_code) == code_norm:
            return fi
        if code_norm in normalize_text(fi.file_name):
            return fi
    return None


def _replace_in_binary_doc(data: bytes, old: str, new: str) -> bytes:
    """老 .doc 二进制替换（仅新旧字节等长时安全）。"""
    if not old or not new:
        return data
    old_b = old.encode("utf-16-le")
    new_b = new.encode("utf-16-le")
    if len(old_b) != len(new_b):
        return data
    if old_b in data:
        return data.replace(old_b, new_b)
    return data


def _word_para_text(para: Any) -> str:
    return para.Range.Text.replace("\r", "").replace("\x07", "").strip()


def _is_signature_date(text: str) -> bool:
    """落款日期：独立短行，不含正文语句。"""
    t = text.strip()
    if not re.search(r"\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日", t):
        return False
    if len(t) > 32 or any(k in t for k in ("之后", "之前", "受理", "检查", "沟通", "。")):
        return False
    return bool(re.fullmatch(r"[\s\d年月日]+", t)) or len(t) <= 24


def _is_signature_company(text: str) -> bool:
    t = text.strip()
    return bool("有限公司" in t and len(t) <= 36)


def _word_find_replace_all(doc: Any, old: str, new: str) -> int:
    """Word 段落级替换（老 .doc 中 Content.Find 常失效，逐段替换更可靠）。"""
    if not old or old == new:
        return 0
    count = 0
    for i in range(1, doc.Paragraphs.Count + 1):
        rng = doc.Paragraphs(i).Range
        text = rng.Text
        if old not in text:
            continue
        n = text.count(old)
        rng.Text = text.replace(old, new)
        count += n
    for ti in range(1, doc.Tables.Count + 1):
        table = doc.Tables(ti)
        for ri in range(1, table.Rows.Count + 1):
            for ci in range(1, table.Rows(ri).Cells.Count + 1):
                rng = table.Cell(ri, ci).Range
                text = rng.Text.replace("\x07", "")
                if old not in text:
                    continue
                n = text.count(old)
                rng.Text = text.replace(old, new)
                count += n
    return count


def _fix_legacy_doc_signature(doc: Any, info: ExtractedInfo, cfg: dict[str, Any]) -> list[str]:
    """修复老 .doc 落款：企业名称单行、日期与企业名右对齐（仅文末落款区）。"""
    wd_align_right = 2
    manufacturer = _val(info, "manufacturer_name")
    new_date = _signing_date(cfg)
    written: list[str] = []

    total = doc.Paragraphs.Count
    tail: list[tuple[int, str, Any]] = []
    for i in range(total, 0, -1):
        p = doc.Paragraphs(i)
        t = _word_para_text(p)
        if t:
            tail.append((i, t, p))
        if len(tail) >= 4:
            break
    tail.reverse()
    if not tail:
        return written

    date_j: int | None = None
    for j in range(len(tail) - 1, -1, -1):
        if _is_signature_date(tail[j][1]):
            date_j = j
            break
    if date_j is None:
        return written

    company_js: list[int] = []
    j = date_j - 1
    while j >= 0 and len(company_js) < 2:
        t = tail[j][1]
        if _is_signature_company(t) or (company_js and len(t.strip()) <= 12):
            company_js.insert(0, j)
            j -= 1
        else:
            break

    if company_js:
        parts = [tail[k][1] for k in company_js]
        joined = re.sub(r"\s+", "", "".join(parts))
        full = manufacturer if manufacturer else joined
        if manufacturer and manufacturer.replace(" ", "") not in joined.replace(" ", ""):
            full = manufacturer
        elif not manufacturer and joined:
            full = joined

        _i, _t, first_p = tail[company_js[0]]
        first_p.Range.Text = full + "\r"
        first_p.Alignment = wd_align_right
        for k in company_js[1:]:
            tail[k][2].Range.Text = "\r"
        written.append(f"落款·企业名称 ← {_clip(full)}")

    _i, _t, date_p = tail[date_j]
    if cfg.get("replace_signing_date", True):
        date_p.Range.Text = new_date + "\r"
    date_p.Alignment = wd_align_right
    written.append(f"落款·日期 ← {new_date if cfg.get('replace_signing_date', True) else _t}")

    return written


def _replace_text_in_word_doc(doc: Any, info: ExtractedInfo, cfg: dict[str, Any]) -> tuple[list[str], list[str]]:
    """Word Find/Replace 逐条替换正文，保留段落/对齐格式（CH1.9 等老 .doc）。"""
    mapping = load_json("field_mapping.json")
    product = _val(info, "product_name")
    patterns = cfg.get("mock_patterns") or mapping.get("mock_product_patterns", [])
    short_patterns = mapping.get("mock_short_patterns", [])
    written: list[str] = []
    skipped: list[str] = []

    replaced_any = False
    body = doc.Content.Text

    if product and cfg.get("replace_mock", True):
        for pat in sorted(patterns, key=len, reverse=True):
            if pat in body:
                n = _word_find_replace_all(doc, pat, product)
                if n:
                    written.append(f"替换「{_clip(pat, 28)}」×{n} → {_clip(product)}")
                    replaced_any = True
                    body = doc.Content.Text
        short = _detection_short_name(info)
        for pat in short_patterns:
            if pat in body:
                n = _word_find_replace_all(doc, pat, short)
                if n:
                    written.append(f"短名替换「{pat}」→ {short}")
                    replaced_any = True
                    body = doc.Content.Text
        if not replaced_any and product:
            skipped.append("正文中未找到模拟产品名占位文本")

    if cfg.get("replace_signing_date", True):
        new_date = _signing_date(cfg)
        for old in _template_signing_dates(cfg):
            if not old or old == new_date:
                continue
            if old in doc.Content.Text:
                n = _word_find_replace_all(doc, old, new_date)
                if n:
                    written.append(f"签署日期 {old} → {new_date}（×{n}）")

    written.extend(_fix_legacy_doc_signature(doc, info, cfg))
    return written, skipped


def _fill_doc_via_word_com(
    template: Path,
    output: Path,
    info: ExtractedInfo,
    cfg: dict[str, Any],
) -> tuple[list[str], list[str], bool]:
    try:
        import win32com.client
    except ImportError:
        return [], ["未安装 pywin32，无法通过 Word 填写 .doc"], False

    written: list[str] = []
    skipped: list[str] = []

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    doc = None
    try:
        doc = word.Documents.Open(str(template.resolve()), ReadOnly=False)
        w, s = _replace_text_in_word_doc(doc, info, cfg)
        written.extend(w)
        skipped.extend(s)
        ensure_dir(output.parent)
        doc.SaveAs2(str(output.resolve()), FileFormat=0)
        doc.Close(False)
        doc = None
        return written, skipped, True
    except Exception as exc:
        skipped.append(f"Word COM 失败：{exc}")
        return written, skipped, False
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        try:
            word.Quit()
        except Exception:
            pass


def _cover_end_index(doc: Document) -> int:
    for i, para in enumerate(doc.paragraphs):
        if "填表说明" in para.text:
            return i
    return 30


def _fill_ch14_cover(
    doc: Document,
    info: ExtractedInfo,
    cfg: dict[str, Any],
) -> tuple[list[str], list[str]]:
    """CH1.4 封面：产品名称、申请人（注册人）。须同时处理段落与表格，避免封面残留模板旧名。"""
    written: list[str] = []
    skipped: list[str] = []
    cover_end = _cover_end_index(doc)
    product = _val(info, "product_name")
    manufacturer = _val(info, "manufacturer_name")
    mapping = load_json("field_mapping.json")
    mock_patterns = cfg.get("mock_patterns") or mapping.get("mock_product_patterns", [])

    if product:
        name_filled = False
        for i, para in enumerate(doc.paragraphs[:cover_end]):
            t = para.text.strip()
            if not t or t in ("申请人：", "申请人"):
                continue
            compact = re.sub(r"\s+", "", t)
            is_product_line = (
                "产品名称" in t
                or any(p in t for p in mock_patterns)
                or any(p.replace(" ", "") in compact for p in mock_patterns)
                or "呼吸道" in t
                or "肺炎支" in t
                or ("核酸检测试剂盒" in compact and i < 25)
            )
            if is_product_line:
                para.text = f"产品名称：{product}"
                written.append(f"封面·产品名称 ← {_clip(product)}")
                name_filled = True
                for j in range(i + 1, min(i + 4, cover_end)):
                    nxt = doc.paragraphs[j].text.strip()
                    if not nxt or nxt in ("申请人：", "申请人"):
                        break
                    if _is_cover_product_fragment(nxt):
                        doc.paragraphs[j].text = ""
                        written.append(f"封面·清除旧名残留段 ← {_clip(nxt, 24)}")
                break
        if not name_filled:
            skipped.append("封面·产品名称（未定位到封面段落）")
    else:
        skipped.append("封面·产品名称（说明书未提取到）")

    if manufacturer:
        for i, para in enumerate(doc.paragraphs[:cover_end]):
            label = para.text.strip().replace("：", "").replace(":", "")
            if label == "申请人":
                for j in range(i + 1, min(i + 4, cover_end)):
                    cand = doc.paragraphs[j].text.strip()
                    if not cand or "产品名称" in cand:
                        continue
                    if cand != manufacturer:
                        doc.paragraphs[j].text = manufacturer
                    written.append(f"封面·申请人 ← {_clip(manufacturer)}")
                    break
                break
    else:
        skipped.append("封面·申请人（说明书未提取注册人名称）")

    if product and mock_patterns:
        for pat in sorted(mock_patterns, key=len, reverse=True):
            n = _replace_in_paragraph_range(doc, pat, product, end=cover_end)
            if n:
                written.append(f"封面·段落替换「{_clip(pat, 24)}」×{n}")
        short = _detection_short_name(info)
        for pat in mapping.get("mock_short_patterns", []):
            n = _replace_in_paragraph_range(doc, pat, short, end=cover_end)
            if n:
                written.append(f"封面·短名替换「{pat}」×{n}")

    _replace_mock_product(doc, info, cfg)
    return written, skipped


def _infer_classification_code(info: ExtractedInfo) -> str:
    name = _val(info, "product_name")
    if any(k in name for k in ("2019-nCoV", "新型冠状病毒", "新冠", "SARS-CoV-2")):
        return "6840-3-017"
    if any(k in name for k in ("合胞", "RSV", "支原体", "MP")):
        return "6840-23-002"
    return "6840-23-002"


def _infer_ch14_profile(info: ExtractedInfo) -> dict[str, Any]:
    name = _val(info, "product_name")
    combined = " ".join(
        _val(info, k)
        for k in (
            "product_name",
            "intended_use",
            "lod",
            "pos_rate",
            "neg_rate",
            "clinical_eval_summary",
        )
    )
    is_covid = any(k in name for k in ("2019-nCoV", "新型冠状病毒", "新冠", "SARS-CoV-2"))
    has_clinical_trial = "临床试验" in _val(info, "clinical_eval_summary")
    has_national_ref = is_covid or bool(re.search(r"国家参考品|国家标准品", combined))
    return {
        "product_class": "第三类",
        "classification_code": _infer_classification_code(info),
        "innovation": "否",
        "main_document": "否",
        "mah_system": "否",
        "national_standard_ref": "是" if has_national_ref else "否",
        "mandatory_national_std": "否",
        "mandatory_industry_std": "否",
        "is_standard_product": "否",
        "clinical_path": "临床试验" if has_clinical_trial else "免于进行临床试验",
        "self_production": True,
        "contract_production": False,
    }


def _table_row_by_label(table: Any, label: str) -> Any | None:
    label_norm = _norm(label)
    for row in table.rows:
        for cell in row.cells:
            if label_norm in _norm(cell.text):
                return row
    return None


def _unique_row_cells(row: Any, start_col: int = 0) -> list[tuple[int, Any]]:
    """合并单元格行内去重，避免同一内容写入多列。"""
    if not row or not row.cells:
        return []
    seen: set[int] = set()
    out: list[tuple[int, Any]] = []
    for i in range(start_col, len(row.cells)):
        tc_id = id(row.cells[i]._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        out.append((i, row.cells[i]))
    return out


def _set_row_value(row: Any, value: str, start_col: int = 1) -> None:
    """仅写入合并行中第一个有效值单元格。"""
    unique = _unique_row_cells(row, start_col)
    if unique:
        unique[0][1].text = value


def _set_merged_row_text(row: Any, value: str, start_col: int = 1) -> None:
    _set_row_value(row, value, start_col)


def _mark_yes_no(choice: str, *, middle: str = "") -> str:
    if choice == "是":
        return f"☑是  {middle}      □否".rstrip()
    return f"□是  {middle}      ☑否".rstrip()


def _mark_applicable(applicable: bool) -> str:
    return "☑适用 □不适用" if applicable else "□适用 ☑不适用"


def _first_phone(contact: str) -> str:
    if not contact:
        return ""
    m = re.search(r"1[3-9]\d{9}|0\d{2,3}-?\d{7,8}", contact)
    if m:
        return m.group(0)
    return contact.split(",")[0].strip()


def _extract_clinical_sites(text: str) -> str:
    """从临床评价摘要中提取机构名称；纯统计描述则留空。"""
    if not text:
        return ""
    sites = re.findall(
        r"[\u4e00-\u9fffA-Za-z0-9（）()·]+(?:医院|疾控中心|检验所|医学科学院|临床检验中心|中心)",
        text,
    )
    deduped = list(dict.fromkeys(s.strip() for s in sites if len(s.strip()) >= 4))
    if deduped:
        return "、".join(deduped)
    if re.search(r"符合率|例样本|总符合", text):
        return ""
    return text.strip() if len(text) <= 80 and "、" in text else ""


def _format_pack_specs_for_ch14(pack_specs: str) -> str:
    """包装规格按规格分行，便于表格单元格阅读。"""
    if not pack_specs:
        return ""
    lines: list[str] = []
    for spec in ("规格A", "规格B"):
        m = re.search(rf"{spec}[：:](.*?)(?=规格[AB]|$)", pack_specs, re.S)
        if not m:
            continue
        block = m.group(1).strip()
        lines.append(f"{spec}：")
        bulk_m = re.search(r"大包装[：:]([^；;\n]+)", block)
        tube_m = re.search(r"分管包装[：:]([^；;\n]+)", block)
        if bulk_m:
            lines.append(f"大包装：{bulk_m.group(1).strip()}；")
        if tube_m:
            lines.append(f"分管包装：{tube_m.group(1).strip().rstrip('。；;')}；")
        if not bulk_m and not tube_m:
            lines.append(block.rstrip("。；;") + "；")
    return "\n".join(lines) if lines else pack_specs


def _format_components_for_ch14(data: Any) -> str:
    parts: list[str] = []
    for spec_label, comps in (("规格A", data.spec_a), ("规格B", data.spec_b)):
        if not comps:
            continue
        parts.append(spec_label)
        for comp in comps:
            line = comp.name
            if comp.composition:
                line = f"{comp.name}：{comp.composition}"
            parts.append(line)
    return "\n".join(parts)


def _set_applicant_table_field(table: Any, label: str, value: str) -> bool:
    label_norm = _norm(label)
    for row in table.rows:
        row_text = "".join(c.text for c in row.cells)
        if "申请人" not in row_text:
            continue
        for i, cell in enumerate(row.cells):
            if _norm(cell.text) == label_norm:
                unique = _unique_row_cells(row, i + 1)
                if unique:
                    unique[0][1].text = value
                    return True
    return False


def _fill_ch14_applicant_section(
    doc: Document,
    info: ExtractedInfo,
) -> tuple[list[str], list[str]]:
    if len(doc.tables) < 2:
        return [], []
    t1 = doc.tables[1]
    written: list[str] = []
    skipped: list[str] = []
    for label, source, log_name in (
        ("名称", "manufacturer_name", "申请人·名称"),
        ("住所", "manufacturer_address", "申请人·住所"),
    ):
        val = _val(info, source)
        if not val:
            skipped.append(f"{log_name}（源数据为空）")
            continue
        if _set_applicant_table_field(t1, label, val):
            written.append(f"{log_name} ← {_clip(val)}")
        else:
            skipped.append(f"{log_name}（未找到表格行）")
    return written, skipped


def _clear_ch14_attachment_placeholders(doc: Document) -> list[str]:
    written: list[str] = []
    for table in doc.tables[:2]:
        for row in table.rows:
            for cell in row.cells:
                if "见附件" in cell.text:
                    cell.text = ""
                    written.append("清除「见附件…」占位")
    return written


def _fill_ch14_attachment_list(
    doc: Document,
    file_list: list[FileInfo] | None,
) -> tuple[list[str], list[str]]:
    if not file_list or len(doc.tables) < 2:
        return [], []
    from core.materials_list import build_ch115_material_lines

    lines = build_ch115_material_lines(file_list)
    if not lines:
        return [], ["应附资料（未识别到可列出的文件）"]
    t1 = doc.tables[1]
    start_i: int | None = None
    for i, row in enumerate(t1.rows):
        if any("应附资料" in c.text for c in row.cells):
            start_i = i
            break
    if start_i is None:
        return [], ["应附资料（未定位表格块）"]
    written: list[str] = []
    data_row_count = len(t1.rows) - start_i - 1
    for offset, line in enumerate(lines):
        text = re.sub(r"^\d+\.\s*", "", line).rstrip("；;")
        if offset < data_row_count:
            _set_row_value(t1.rows[start_i + 1 + offset], text, 0)
            written.append(f"应附资料·{offset + 1} ← {_clip(line, 36)}")
        else:
            last_row = t1.rows[start_i + data_row_count]
            unique = _unique_row_cells(last_row, 0)
            prev = unique[0][1].text.strip() if unique else ""
            merged = f"{prev}\n{text}" if prev else text
            _set_row_value(last_row, merged, 0)
            written.append(f"应附资料·{offset + 1} ← {_clip(line, 36)}（续）")
    for row_i in range(start_i + 1 + min(len(lines), data_row_count), start_i + 1 + data_row_count):
        if row_i < len(t1.rows):
            _set_row_value(t1.rows[row_i], "", 0)
    return written, []


def _cleanup_ch14_trailing_content(doc: Document) -> list[str]:
    written: list[str] = []
    markers = ("器审中心", "合规要求提示", "本次不勾选", "受贿行贿", "亲清新型政商关系")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if any(m in cell.text for m in markers):
                    cell.text = ""
                    if "删除合规提示" not in written:
                        written.append("删除器审中心合规提示等非标准内容")
    if len(doc.tables) >= 3:
        for cell in doc.tables[2].rows[0].cells:
            if cell.text.strip():
                cell.text = ""
                written.append("清除合规提示续页")
    if len(doc.tables) >= 3:
        t2 = doc.tables[2]
        if len(t2.rows) > 2:
            guarantee_row = t2.rows[2]
            gtext = "".join(c.text for c in guarantee_row.cells)
            if "保证书" in gtext:
                unique = _unique_row_cells(guarantee_row, 0)
                if len(unique) > 1:
                    for _, cell in unique[1:]:
                        cell.text = ""
                    written.append("保证书去重（仅保留首列）")
    return written


def _manual_path_from_files(file_list: list[FileInfo] | None) -> Path | None:
    if not file_list:
        return None
    for fi in file_list:
        if fi.file_category == "核心数据源" and "说明书" in fi.file_name:
            return Path(fi.file_path)
    return None


def _fill_ch14_derived_fields(
    doc: Document,
    info: ExtractedInfo,
    *,
    file_list: list[FileInfo] | None,
) -> tuple[list[str], list[str]]:
    written: list[str] = []
    skipped: list[str] = []
    if len(doc.tables) < 1:
        return written, skipped
    t0 = doc.tables[0]

    pack = _format_pack_specs_for_ch14(_val(info, "pack_specs"))
    if pack:
        row = _table_row_by_label(t0, "包装规格")
        if row:
            _set_row_value(row, pack, 1)
            written.append(f"包装规格（分行） ← {_clip(pack, 40)}")

    manual = _manual_path_from_files(file_list)
    components_text = ""
    if manual and manual.exists():
        try:
            from core.product_list import extract_product_list_by_rules

            pl_data = extract_product_list_by_rules(manual, info)
            components_text = _format_components_for_ch14(pl_data)
        except Exception:
            components_text = ""
    if components_text:
        row = _table_row_by_label(t0, "主要组成成分")
        if row:
            _set_row_value(row, components_text[:800], 1)
            written.append(f"主要组成成分 ← {_clip(components_text, 40)}")
    else:
        skipped.append("主要组成成分（说明书未解析到组分表）")

    udi_row = _table_row_by_label(t0, "UDI-DI")
    if udi_row:
        _set_row_value(udi_row, "", 1)

    return written, skipped


def _replace_ch14_table_signing_date(doc: Document, cfg: dict[str, Any]) -> list[str]:
    new_date = _signing_date(cfg)
    m = re.match(r"(\d{4})年(\d{1,2})月(\d{1,2})日", new_date)
    if not m:
        return []
    year, month, day = m.group(1), m.group(2).zfill(2), m.group(3).zfill(2)
    written: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_text = "".join(c.text for c in row.cells)
            if "填表人" not in row_text or "日期" not in row_text:
                continue
            for cell in row.cells:
                if re.search(r"\d{4}\s*年", cell.text):
                    cell.text = f"  填表人： 日期： {year} 年 {month} 月"
                    written.append(f"填表日期 → {new_date}")
                elif re.fullmatch(r"\s*\d{1,2}\s*日\s*", cell.text.strip()) or (
                    "日" in cell.text and "年" not in cell.text and len(cell.text.strip()) <= 8
                ):
                    cell.text = f"  {day} 日"
    return written


def _fill_ch14_form_options(
    doc: Document,
    info: ExtractedInfo,
    cfg: dict[str, Any],
) -> tuple[list[str], list[str]]:
    """CH1.4 分类编码、勾选项、临床路径、生产形式等。"""
    if len(doc.tables) < 2:
        return [], ["CH1.4 表格结构异常，跳过表单选项填写"]

    profile = _infer_ch14_profile(info)
    overrides = cfg.get("form_options") or {}
    profile.update({k: v for k, v in overrides.items() if v is not None})

    written: list[str] = []
    skipped: list[str] = []
    t0 = doc.tables[0]
    t1 = doc.tables[1]

    def set_labeled(label: str, formatter, field_name: str) -> None:
        row = _table_row_by_label(t0, label)
        if not row:
            skipped.append(f"{field_name}（未找到「{label}」行）")
            return
        val = formatter() if callable(formatter) else formatter
        _set_merged_row_text(row, val)
        written.append(f"{field_name} ← {_clip(str(val), 40)}")

    set_labeled("产品类别", lambda: f"☑{profile['product_class']} □同第三类", "产品类别")
    set_labeled("创新产品", lambda: _mark_yes_no(profile["innovation"], middle="通过创新审查申请受理号："), "创新产品")
    set_labeled("涉及主文档", lambda: _mark_yes_no(profile["main_document"], middle="主文档登记号："), "涉及主文档")
    set_labeled("注册人制度", lambda: _mark_yes_no(profile["mah_system"], middle="企业名称："), "注册人制度")
    set_labeled(
        "优先通道",
        "□优先通道申请   □应急通道    □同品种首个产品首次申报   □真实世界数据应用试点",
        "优先通道",
    )
    set_labeled("分类编码", profile["classification_code"], "分类编码")
    set_labeled(
        "适用的国家标准品",
        lambda: _mark_yes_no(profile["national_standard_ref"]),
        "国家标准品",
    )
    set_labeled(
        "强制性国家标准",
        lambda: _mark_yes_no(profile["mandatory_national_std"]),
        "强制性国家标准",
    )
    set_labeled(
        "强制性行业标准",
        lambda: _mark_yes_no(profile["mandatory_industry_std"]),
        "强制性行业标准",
    )

    std_row = _table_row_by_label(t0, "是否属于标准品")
    if std_row:
        _set_row_value(std_row, _mark_yes_no(profile["is_standard_product"]), 1)
        written.append(f"标准品/参考品 ← {profile['is_standard_product']}")
    else:
        skipped.append("标准品/参考品（未找到对应行）")

    clinical_raw = _val(info, "clinical_eval_summary").lstrip("：: ")
    if profile["clinical_path"] == "临床试验":
        sites = _extract_clinical_sites(clinical_raw)
        clinical_text = "□免于进行临床试验 ☑临床试验"
        if sites:
            clinical_text += f"\n临床试验机构名称：{sites}"
        else:
            clinical_text += "\n临床试验机构名称："
    else:
        clinical_text = "☑免于进行临床试验 □临床试验"
    clinical_row = _table_row_by_label(t1, "临床评价路径")
    if clinical_row:
        _set_row_value(clinical_row, clinical_text, 1)
        written.append(f"临床评价路径 ← {_clip(clinical_text, 50)}")
    else:
        skipped.append("临床评价路径（未找到对应行）")

    addr = _val(info, "production_address") or _val(info, "manufacturer_address")
    self_row = _table_row_by_label(t1, "申请人自行生产")
    if self_row:
        if profile["self_production"]:
            self_text = "申请人自行生产：☑适用 □不适用"
            if addr:
                self_text += f"\n{addr}"
        else:
            self_text = "申请人自行生产：□适用 ☑不适用"
        _set_row_value(self_row, self_text, 1)
        written.append(f"自行生产 ← {'适用' if profile['self_production'] else '不适用'}")
    contract_row = _table_row_by_label(t1, "委托生产")
    if contract_row:
        contract_text = (
            "非申请人自行生产（含委托生产）：☑适用 □不适用"
            if profile["contract_production"]
            else "非申请人自行生产（含委托生产）：□适用 ☑不适用"
        )
        _set_row_value(contract_row, contract_text, 1)
        written.append(f"委托生产 ← {'适用' if profile['contract_production'] else '不适用'}")

    phone = _first_phone(_val(info, "contact_info"))
    if phone:
        for row in t1.rows:
            row_text = "".join(c.text for c in row.cells)
            if "申请人" not in row_text or "手机号" not in row_text:
                continue
            for i, cell in enumerate(row.cells):
                if "手机号" in cell.text:
                    unique = _unique_row_cells(row, i + 1)
                    if unique:
                        unique[0][1].text = phone
                        written.append(f"联系人·手机号 ← {phone}")
                    break
            break

    return written, skipped


def fill_table_fields(
    template: Path,
    output: Path,
    info: ExtractedInfo,
    cfg: dict[str, Any],
    *,
    file_list: list[FileInfo] | None = None,
) -> Path:
    ensure_dir(output.parent)
    shutil.copy2(template, output)
    doc = Document(str(output))
    written: list[str] = []
    skipped: list[str] = []

    if cfg.get("fill_cover", True):
        cover_written, cover_skipped = _fill_ch14_cover(doc, info, cfg)
        written.extend(cover_written)
        skipped.extend(cover_skipped)

    for field in cfg.get("fields", []):
        val = _val(info, field["source"])
        label = field["row_label"]
        if field["source"] == "pack_specs" and val:
            val = _format_pack_specs_for_ch14(val)
        if not val and field.get("optional"):
            skipped.append(f"{label}（源数据为空，可选项已跳过）")
            continue
        if not val:
            skipped.append(f"{label}（说明书里没提到 {field['source']}）")
            continue
        if _set_table_cell_by_label(doc, label, val, max_len=field.get("max_len")):
            written.append(f"{label} ← {_clip(val)}")
        else:
            skipped.append(f"{label}（表格里找不到这一行）")

    if cfg.get("fill_form_options", True):
        opt_written, opt_skipped = _fill_ch14_form_options(doc, info, cfg)
        written.extend(opt_written)
        skipped.extend(opt_skipped)

    app_written, app_skipped = _fill_ch14_applicant_section(doc, info)
    written.extend(app_written)
    skipped.extend(app_skipped)

    if cfg.get("fill_derived_fields", True):
        derived_written, derived_skipped = _fill_ch14_derived_fields(
            doc, info, file_list=file_list
        )
        written.extend(derived_written)
        skipped.extend(derived_skipped)

    if cfg.get("clear_attachment_placeholders", True):
        written.extend(_clear_ch14_attachment_placeholders(doc))

    if cfg.get("sync_attachment_list", True) and file_list:
        att_written, att_skipped = _fill_ch14_attachment_list(doc, file_list)
        written.extend(att_written)
        skipped.extend(att_skipped)

    if cfg.get("cleanup_trailing_content", True):
        written.extend(_cleanup_ch14_trailing_content(doc))

    if cfg.get("replace_signing_date", True):
        written.extend(_replace_signing_dates_docx(doc, cfg))
        written.extend(_replace_ch14_table_signing_date(doc, cfg))

    doc.save(str(output))
    _log_doc_fill(
        output,
        template,
        "表格字段 + 封面（CH1.4 申请表）",
        written=written,
        skipped=skipped,
    )
    return output


def fill_declaration(
    template: Path,
    output: Path,
    info: ExtractedInfo,
    cfg: dict[str, Any],
    *,
    file_list: list[FileInfo] | None = None,
    filled_paths: list[str] | None = None,
) -> Path:
    ensure_dir(output.parent)
    shutil.copy2(template, output)
    doc = Document(str(output))
    product = _val(info, "product_name")
    manufacturer = _val(info, "manufacturer_name")
    written: list[str] = []
    skipped: list[str] = []

    pattern = cfg.get(
        "declaration_pattern",
        r"申请境内第三类体外诊断试剂(.+?)(产品注册|注册)",
    )

    replaced = False
    if product:
        for para in doc.paragraphs:
            text = para.text
            if "申请境内第三类" in text:
                new_text = re.sub(
                    pattern,
                    lambda m: f"申请境内第三类体外诊断试剂{product}{m.group(2) if m.lastindex and m.lastindex >= 2 else '注册'}",
                    text,
                    count=1,
                )
                if new_text != text:
                    para.text = new_text
                    replaced = True
            elif manufacturer and "我公司" in text and "申请" in text:
                new_text = re.sub(
                    r"申请境内第三类体外诊断试剂(.+?)(产品注册|注册)",
                    f"申请境内第三类体外诊断试剂{product}\\2",
                    text,
                    count=1,
                )
                if new_text != text:
                    para.text = new_text
                    replaced = True
        if replaced:
            written.append(f"声明正文产品名称 ← {_clip(product)}")
        else:
            skipped.append("声明段落（未找到「申请境内第三类…」句式）")
    else:
        skipped.append("产品名称（说明书未提取到，无法替换声明）")

    _replace_mock_product(doc, info, cfg)
    if product:
        written.append("模拟产品名全文替换（旧模板 RSV/MP → 说明书产品名）")

    if cfg.get("sync_materials_list") and file_list:
        from core.materials_list import build_ch115_material_lines

        material_lines = build_ch115_material_lines(file_list, filled_paths=filled_paths)
        if material_lines and _rewrite_paragraph_block(
            doc,
            "提交如下材料：",
            "我单位保证",
            material_lines,
        ):
            written.append(f"材料清单 ← {len(material_lines)} 项（按实际上传/已填写文件）")
        elif not material_lines:
            skipped.append("材料清单（未从上传文件夹识别到可列出的申报材料）")
        else:
            skipped.append("材料清单（未定位到「提交如下材料：」段落块）")

    if cfg.get("replace_signing_date", True):
        written.extend(_replace_signing_dates_docx(doc, cfg))

    doc.save(str(output))
    _log_doc_fill(output, template, "声明类文档（正则替换产品名）", written=written, skipped=skipped)
    return output


def _para_index(doc: Document, needle: str) -> int | None:
    for i, para in enumerate(doc.paragraphs):
        if needle in para.text:
            return i
    return None


def _insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.text = text
    return new_para


def _format_standard_line(entry: Any) -> str:
    return f"{entry.std_no}    {entry.name}"


def _rewrite_paragraph_block(
    doc: Document,
    start_marker: str,
    end_marker: str,
    lines: list[str],
    *,
    drop_start_heading: bool = False,
) -> bool:
    start_i = _para_index(doc, start_marker)
    end_i = _para_index(doc, end_marker)
    if start_i is None or end_i is None or end_i <= start_i:
        return False

    content_start = start_i + 1
    if drop_start_heading:
        doc.paragraphs[start_i].text = ""
        content_start = start_i

    for i in range(content_start, end_i):
        doc.paragraphs[i].text = ""

    anchor = doc.paragraphs[start_i]
    for j, line in enumerate(lines):
        target_i = content_start + j
        if target_i < end_i:
            doc.paragraphs[target_i].text = line
            anchor = doc.paragraphs[target_i]
        else:
            anchor = _insert_paragraph_after(anchor, line)
    return True


def fill_standards_list(
    template: Path,
    output: Path,
    info: ExtractedInfo,
    cfg: dict[str, Any],
    *,
    file_list: list[FileInfo] | None = None,
) -> Path:
    del file_list
    ensure_dir(output.parent)
    shutil.copy2(template, output)
    doc = Document(str(output))
    product = _val(info, "product_name")
    manufacturer = _val(info, "manufacturer_name")
    written: list[str] = []
    skipped: list[str] = []

    from core.standards import resolve_standards

    std_data = resolve_standards(info, use_llm=cfg.get("use_llm", True))

    # 声明正文
    decl_idx = _para_index(doc, "申请境内第三类")
    if decl_idx is not None and product:
        para = doc.paragraphs[decl_idx]
        old = para.text
        mfr = manufacturer or "本公司"
        new_text = re.sub(
            r"^(.+?)申请境内第三类体外诊断试剂(.+?)(产品注册|注册)",
            rf"\1申请境内第三类体外诊断试剂{product}\3",
            old,
            count=1,
        )
        if new_text == old:
            new_text = re.sub(
                r"申请境内第三类体外诊断试剂(.+?)(产品注册|注册)",
                rf"申请境内第三类体外诊断试剂{product}\2",
                old,
                count=1,
            )
        if manufacturer and manufacturer not in new_text:
            new_text = re.sub(r"^[^\s，,]+(?:公司|有限公司)", manufacturer, new_text, count=1)
        para.text = new_text
        written.append(f"声明正文 ← 产品名称 {_clip(product)}")
    elif not product:
        skipped.append("声明正文（未提取到产品名称）")

    # 适用标准清单段落
    lines: list[str] = [_format_standard_line(s) for s in std_data.standards]
    if std_data.guidance:
        lines.append("")
        lines.append("【产品相关技术指南 / 指导原则（本地库匹配，V1.0 可扩展 CMDE 在线检索）】")
        for g in std_data.guidance:
            line = f"· {g.title}（{g.source}）"
            if g.note:
                line += f" — {g.note}"
            lines.append(line)
    lines.append("")
    lines.append(f"【系统自动匹配说明】{std_data.match_summary}数据来源：{std_data.source}。")

    title_i = _para_index(doc, "符合标准的清单")
    start_i = _para_index(doc, "符合的标准清单如下：")
    drop_heading = title_i is not None and start_i is not None and title_i < start_i

    if _rewrite_paragraph_block(
        doc,
        "符合的标准清单如下：",
        "符合的标准品清单如下：",
        lines,
        drop_start_heading=drop_heading,
    ):
        written.append(
            f"适用标准清单 ← {len(std_data.standards)} 项标准 + {len(std_data.guidance)} 项指导原则"
            + ("（已合并重复标题行）" if drop_heading else "")
        )
    else:
        skipped.append("适用标准清单（未定位到「符合的标准清单如下：」段落）")

    std_prod_start = _para_index(doc, "符合的标准品清单如下：")
    if std_prod_start is not None and std_prod_start + 1 < len(doc.paragraphs):
        prod_para = doc.paragraphs[std_prod_start + 1]
        if std_data.standard_products:
            prod_para.text = "\n".join(f"· {p}" for p in std_data.standard_products)
            written.append(f"标准品清单 ← {len(std_data.standard_products)} 项")
        else:
            prod_para.text = "无"
            if std_data.standard_products_note:
                note_idx = std_prod_start + 2
                if note_idx < len(doc.paragraphs) and not doc.paragraphs[note_idx].text.strip():
                    doc.paragraphs[note_idx].text = std_data.standard_products_note
                else:
                    _insert_paragraph_after(prod_para, std_data.standard_products_note)
            written.append("标准品清单 ← 无（已输出系统检索说明）")

    # 签署日期 / 申请企业
    if manufacturer:
        for para in doc.paragraphs:
            if para.text.strip() in ("卡尤迪生物科技宜兴有限公司",) or (
                "有限公司" in para.text and len(para.text.strip()) < 30 and "申请" not in para.text
            ):
                if manufacturer not in para.text:
                    para.text = manufacturer
                    written.append(f"落款企业 ← {_clip(manufacturer)}")
                    break

    if cfg.get("replace_signing_date", True):
        written.extend(_replace_signing_dates_docx(doc, cfg))

    doc.save(str(output))
    _log_doc_fill(
        output,
        template,
        "CH1.11.1 标准清单（本地库+说明书匹配+可选LLM → Python 写段落）",
        written=written,
        skipped=skipped,
    )
    return output


def fill_product_list(
    template: Path,
    output: Path,
    info: ExtractedInfo,
    cfg: dict[str, Any],
    *,
    file_list: list[FileInfo] | None = None,
) -> Path:
    del file_list
    ensure_dir(output.parent)
    shutil.copy2(template, output)
    doc = Document(str(output))
    written: list[str] = []
    skipped: list[str] = []

    product = _val(info, "product_name")
    pack = _val(info, "pack_specs")

    if product:
        for para in doc.paragraphs:
            if "CH1.5" in para.text or "产品列表" in para.text:
                continue
            t = para.text.strip()
            if t.endswith("）") and len(t) < 40 and "PCR" in t:
                continue
            if "呼吸道" in t or ("核酸检测试剂盒" in t and len(t) < 80):
                para.text = product.split("（")[0].strip() if "（" in product else product
            if "包装规格" in para.text or "的包装规格、货号" in para.text:
                para.text = f"{product}的包装规格、货号、规格及主要组成成分如下表"
        written.append(f"标题/引导段 ← {_clip(product)}")

    from core.product_list import apply_product_list_to_doc

    table_written, table_skipped = apply_product_list_to_doc(doc, info)
    written.extend(table_written)
    skipped.extend(table_skipped)

    if cfg.get("replace_signing_date", True):
        written.extend(_replace_signing_dates_docx(doc, cfg))

    doc.save(str(output))
    _log_doc_fill(
        output,
        template,
        "产品列表 CH1.5（说明书组分表 + 大模型 → Python 写表）",
        written=written,
        skipped=skipped,
    )
    return output


def fill_ch12_catalog(
    template: Path,
    output: Path,
    info: ExtractedInfo,
    cfg: dict[str, Any],
    *,
    file_list: list[FileInfo] | None = None,
) -> Path:
    """CH1.2 监管信息目录：更新封面产品名，并按扫描结果同步页码。"""
    ensure_dir(output.parent)
    shutil.copy2(template, output)
    doc = Document(str(output))
    written: list[str] = []
    skipped: list[str] = []

    if cfg.get("replace_mock", True):
        _replace_mock_product(doc, info, cfg)
        if _val(info, "product_name"):
            written.append(f"目录封面产品名 ← {_clip(_val(info, 'product_name'))}")

    product = _val(info, "product_name")
    if product:
        for para in doc.paragraphs:
            if "注册申报资料" in para.text:
                continue
            if "呼吸道" in para.text or "核酸检测试剂盒" in para.text:
                para.text = product

    synced = 0
    if cfg.get("sync_pages", True) and file_list and doc.tables:
        table = doc.tables[0]
        for row in table.rows[1:]:
            cells = row.cells
            if len(cells) < 5:
                continue
            code = cells[0].text.strip()
            if not code.startswith("CH"):
                continue
            matched = _find_file_for_chapter(file_list, code)
            if matched:
                cells[3].text = Path(matched.file_name).name
                cells[4].text = str(matched.page_count)
                synced += 1
        if synced:
            written.append(f"CH1.2 目录表已同步 {synced} 行页码与文件名")
        else:
            skipped.append("页码同步（未匹配到 CH 章节行）")
    elif cfg.get("sync_pages", True):
        skipped.append("页码同步（缺少文件清单或表格）")

    doc.save(str(output))
    _log_doc_fill(output, template, "章节目录 CH1.2（封面 + 页码同步）", written=written, skipped=skipped)
    return output


def fill_legacy_doc(
    template: Path,
    output: Path,
    info: ExtractedInfo,
    cfg: dict[str, Any],
    *,
    file_list: list[FileInfo] | None = None,
) -> Path:
    """CH1.9 等老 .doc：优先 Word COM 替换；失败则仅复制模版（避免二进制替换损坏文件）。"""
    del file_list
    ensure_dir(output.parent)
    product = _val(info, "product_name")

    written, skipped, ok = _fill_doc_via_word_com(template, output, info, cfg)
    if ok:
        _log_doc_fill(
            output,
            template,
            "Word COM 填写老 .doc（CH1.9 等）",
            written=written,
            skipped=skipped,
            note="通过 Word Find/Replace 替换正文并修复落款右对齐，避免 Content.Text 破坏排版",
        )
        return output

    shutil.copy2(template, output)
    skipped.append("已回退为复制模版原文件（请在 Word 中手工改产品名，或安装 Word + pywin32）")
    _log_doc_fill(
        output,
        template,
        "老 .doc 仅复制模版（COM 不可用）",
        written=["已复制模版，正文保持原样"] if output.exists() else [],
        skipped=skipped,
        note="切勿对 .doc 做不等长二进制替换，否则 Word 打开会空白",
    )
    return output


_FILLERS = {
    "table_fields": fill_table_fields,
    "declaration": fill_declaration,
    "standards_list": fill_standards_list,
    "product_list": fill_product_list,
    "ch12_catalog": fill_ch12_catalog,
    "legacy_doc": fill_legacy_doc,
}


def fill_templates(
    upload_dir: Path,
    output_dir: Path,
    info: ExtractedInfo,
    *,
    file_list: list[FileInfo] | None = None,
) -> list[str]:
    mapping_cfg = load_json("field_mapping.json")
    filled: list[str] = []
    targets_cfg = mapping_cfg.get("targets", {})

    for _key, cfg in targets_cfg.items():
        tpl_glob = cfg.get("template_glob")
        if not tpl_glob:
            continue
        templates = glob_files(upload_dir, tpl_glob)
        if not templates:
            continue

        output_name = cfg.get("output", f"{_key}_已填写.docx")
        output_path = output_dir / output_name
        fill_type = cfg.get("fill_type", "table_fields")
        filler = _FILLERS.get(fill_type, fill_table_fields)
        if fill_type == "declaration" and cfg.get("sync_materials_list"):
            filler(
                templates[0],
                output_path,
                info,
                cfg,
                file_list=file_list,
                filled_paths=list(filled),
            )
        else:
            filler(templates[0], output_path, info, cfg, file_list=file_list)
        filled.append(str(output_path))
        stale = _audit_stale_signing_dates(output_path, cfg)
        if stale:
            from core.run_logger import get_run_logger

            logger = get_run_logger()
            if logger:
                logger.bullet(
                    f"⚠ {output_path.name} 仍含旧签署日期 {stale}，请检查填写逻辑"
                )

    return filled
