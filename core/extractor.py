from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document

from core.llm_client import extract_with_llm, should_use_llm
from core.utils import glob_files, load_json, normalize_text


@dataclass
class ExtractedInfo:
    product_name: str = "未提及"
    pack_specs: str = "未提及"
    intended_use: str = "未提及"
    storage_condition: str = "未提及"
    detection_principle: str = "未提及"
    sample_types: str = "未提及"
    instruments: list[str] = field(default_factory=list)
    manufacturer_name: str = "未提及"
    manufacturer_address: str = "未提及"
    production_address: str = "未提及"
    production_license: str = "未提及"
    clinical_eval_summary: str = "未提及"
    contact_info: str = "未提及"
    lod: str = "未提及"
    pos_rate: str = "未提及"
    neg_rate: str = "未提及"
    targets: list[str] = field(default_factory=list)
    product_list: dict[str, Any] = field(default_factory=dict)
    confidence: dict[str, str] = field(default_factory=dict)
    source_file: str = ""
    llm_used: bool = False

    def to_dict(self) -> dict[str, Any]:
        return {
            "product_name": self.product_name,
            "pack_specs": self.pack_specs,
            "intended_use": self.intended_use,
            "storage_condition": self.storage_condition,
            "detection_principle": self.detection_principle,
            "sample_types": self.sample_types,
            "instruments": self.instruments,
            "manufacturer_name": self.manufacturer_name,
            "manufacturer_address": self.manufacturer_address,
            "production_address": self.production_address,
            "production_license": self.production_license,
            "clinical_eval_summary": self.clinical_eval_summary,
            "contact_info": self.contact_info,
            "lod": self.lod,
            "pos_rate": self.pos_rate,
            "neg_rate": self.neg_rate,
            "targets": self.targets,
            "product_list": self.product_list,
            "confidence": self.confidence,
            "source_file": self.source_file,
            "llm_used": self.llm_used,
        }


SECTION_PATTERNS: dict[str, str] = {
    "product_name": r"【产品名称】\s*\n?\s*(.+?)(?=\n【|\Z)",
    "pack_specs": r"【包装规格】\s*\n?\s*(.+?)(?=\n【|\Z)",
    "intended_use": r"【预期用途】\s*\n?\s*(.+?)(?=\n【|\Z)",
    "storage_condition": r"【储存条件及有效期】\s*\n?\s*(.+?)(?=\n【[^】]{0,20}】|\Z)",
    "detection_principle": r"【检测原理】\s*\n?\s*(.+?)(?=\n【|\Z)",
    "sample_types": r"适用样本类型[：:]\s*(.+?)(?=\n|\Z)",
    "instruments": r"【适用仪器】\s*\n?\s*(.+?)(?=\n【|\Z)",
}


def read_docx_full_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


def extract_section(text: str, pattern: str, multiline: bool = True) -> str | None:
    flags = re.S if multiline else 0
    m = re.search(pattern, text, flags)
    if not m:
        return None
    value = m.group(1).strip()
    value = re.sub(r"\s+", " ", value)
    return value[:2000] if value else None


def extract_line(text: str, pattern: str) -> str | None:
    m = re.search(pattern, text)
    return m.group(1).strip() if m else None


def extract_by_rules(text: str) -> ExtractedInfo:
    info = ExtractedInfo()
    conf: dict[str, str] = {}

    for key, pattern in SECTION_PATTERNS.items():
        val = extract_section(text, pattern)
        if val:
            if key == "instruments":
                info.instruments = [x.strip() for x in re.split(r"[、,，]", val) if x.strip()]
            else:
                setattr(info, key, val)
            conf[key] = "高(规则)"

    info.manufacturer_name = extract_line(text, r"注册人/售后服务单位名称[：:]\s*(.+)") or info.manufacturer_name
    if info.manufacturer_name != "未提及":
        conf["manufacturer_name"] = "高(规则)"
    info.manufacturer_address = extract_line(text, r"生产企业住所[：:]\s*(.+)") or info.manufacturer_address
    info.production_address = extract_line(text, r"生产地址[：:]\s*(.+)") or info.production_address
    info.production_license = extract_line(text, r"生产许可证编号[：:]\s*(.+)") or info.production_license
    info.contact_info = extract_line(text, r"联系方式[：:]\s*(.+)") or info.contact_info
    clinical = extract_section(text, r"【?临床评价】?\s*\n?\s*(.+?)(?=\n【|\Z)")
    if not clinical:
        clinical = extract_line(text, r"10\.临床评价[：:]\s*(.+?)(?=\n【|\Z)")
    if clinical:
        info.clinical_eval_summary = clinical[:500]
        conf["clinical_eval_summary"] = "高(规则)"
    info.lod = extract_line(text, r"最低检出限[：:为]*\s*(.+?)(?=\n|。)") or extract_line(text, r"最低检测限[：:为]*\s*(.+?)(?=\n|。)") or info.lod
    info.pos_rate = extract_line(text, r"阳性符合率[：:为]*\s*(.+?)(?=\n|。)") or info.pos_rate
    info.neg_rate = extract_line(text, r"阴性符合率[：:为]*\s*(.+?)(?=\n|。)") or info.neg_rate

    # 从预期用途推断检测靶标（规则）
    targets: list[str] = []
    for m in re.finditer(r"(ORF1ab|N基因|E基因|S基因|RdRp|[A-Za-z0-9]+基因)", text):
        t = m.group(1)
        if t not in targets:
            targets.append(t)
    if targets:
        info.targets = targets
        conf["targets"] = "高(规则)"

    if info.lod != "未提及":
        conf["lod"] = "高(规则)"
    if info.pos_rate != "未提及":
        conf["pos_rate"] = "高(规则)"

    info.confidence = conf
    return info


# 启用 LLM 时，对这些语义字段用模型结果精炼/摘要（规则仍保留作对照）
LLM_ENHANCE_FIELDS = frozenset(
    {
        "intended_use",
        "detection_principle",
        "sample_types",
        "storage_condition",
        "manufacturer_name",
        "manufacturer_address",
        "contact_info",
        "lod",
        "pos_rate",
        "neg_rate",
        "detection_targets",
    }
)


def _trim_intended_use(text: str) -> str:
    """预期用途取首段核心表述，便于填入申请表。"""
    text = re.sub(r"\s+", " ", text.strip())
    for sep in ("。", "；", "\n"):
        if sep in text:
            first = text.split(sep)[0].strip()
            if len(first) > 20:
                return first + ("。" if sep == "。" else "")
    return text[:400]


def merge_llm(info: ExtractedInfo, llm_data: dict[str, Any], *, enhance: bool = False) -> ExtractedInfo:
    for key, val in llm_data.items():
        if key in ("confidence", "source_section", "source_file", "llm_used"):
            continue
        if key == "detection_targets" and isinstance(val, list):
            info.targets = [str(x).strip() for x in val if str(x).strip()]
            if info.targets:
                info.confidence["targets"] = "中(LLM)" if enhance else "中(LLM)"
            continue
        if key == "detection_targets" and isinstance(val, str) and val != "未提及":
            info.targets = [x.strip() for x in re.split(r"[、,，;；]", val) if x.strip()]
            if info.targets:
                info.confidence["targets"] = "中(LLM)"
            continue
        if key in ("targets",):
            continue
        if val in (None, "", "未提及"):
            continue

        if key == "instruments" and isinstance(val, list):
            llm_val: Any = [str(x).strip() for x in val if str(x).strip()]
        elif key == "instruments" and isinstance(val, str):
            llm_val = [x.strip() for x in re.split(r"[、,，;；]", val) if x.strip()]
        elif key == "intended_use" and isinstance(val, str):
            llm_val = _trim_intended_use(val)
        elif key == "detection_principle" and isinstance(val, str):
            llm_val = val[:500]
        else:
            llm_val = val

        current = getattr(info, key, None)

        if enhance and key in LLM_ENHANCE_FIELDS:
            # 储存条件：规则常能抓全段，LLM 摘要会丢冻融/开瓶稳定性，保留更完整的一段
            if key == "storage_condition" and current not in (None, "未提及", "", []):
                cur_s, llm_s = str(current).strip(), str(llm_val).strip()
                if len(cur_s) >= len(llm_s) and cur_s not in ("未提及",):
                    setattr(info, key, cur_s)
                    cur_norm = normalize_text(cur_s)
                    llm_norm = normalize_text(llm_s)
                    info.confidence[key] = (
                        "高(规则+LLM)" if llm_norm and (cur_norm in llm_norm or llm_norm in cur_norm) else "高(规则)"
                    )
                else:
                    setattr(info, key, llm_val)
                    info.confidence[key] = "中(LLM)"
                continue

            setattr(info, key, llm_val)
            cur_norm = normalize_text(str(current)) if current not in (None, "未提及", "", []) else ""
            llm_norm = normalize_text(str(llm_val)) if not isinstance(llm_val, list) else ""
            if cur_norm and llm_norm and (cur_norm in llm_norm or llm_norm in cur_norm or cur_norm == llm_norm):
                info.confidence[key] = "高(规则+LLM)"
            else:
                info.confidence[key] = "中(LLM)"
            continue

        if current in (None, "未提及", [], ""):
            setattr(info, key, llm_val)
            info.confidence[key] = "中(LLM)"

    return info


def find_manual(upload_dir: Path) -> Path | None:
    files = glob_files(upload_dir, "*说明书*.docx")
    return files[0] if files else None


def extract_from_upload(upload_dir: Path, *, use_llm: bool = True) -> ExtractedInfo:
    manual = find_manual(upload_dir)
    if not manual:
        return ExtractedInfo(confidence={"_error": "未找到产品说明书"})

    text = read_docx_full_text(manual)
    info = extract_by_rules(text)
    info.source_file = manual.name

    from core.run_logger import get_run_logger

    logger = get_run_logger()
    if logger:
        rule_fields = [k for k, v in info.confidence.items() if v]
        logger.python_only(
            "阶段3 · 规则正则先扫说明书",
            f"从《{manual.name}》用正则抽出 {len(rule_fields)} 个字段，"
            f"产品名称={info.product_name[:40]}{'…' if len(info.product_name) > 40 else ''}",
        )

    if should_use_llm(use_llm):
        info.llm_used = True
        llm_data = extract_with_llm(text)
        if llm_data:
            info = merge_llm(info, llm_data, enhance=True)
            if logger:
                enhanced = [k for k, v in info.confidence.items() if "LLM" in str(v)]
                name_display = f"{info.product_name[:30]}…" if len(info.product_name) > 30 else info.product_name
                logger.python_only(
                    "阶段3 · Python 合并大模型结果",
                    f"大模型精炼/补充了 {len(enhanced)} 个字段：{', '.join(enhanced[:8]) or '无'}；"
                    f"产品名称仍为「{name_display}」",
                )
        else:
            info.confidence["_llm"] = "调用失败或未返回有效 JSON"
            if logger:
                logger.python_only("阶段3 · 大模型未生效", "全程使用规则提取结果。")

    for field_name in ("product_name", "intended_use", "pack_specs"):
        if field_name not in info.confidence and getattr(info, field_name) != "未提及":
            info.confidence[field_name] = "高(规则)"

    from core.product_list import extract_product_list

    pl = extract_product_list(manual, info, use_llm=use_llm)
    info.product_list = pl.to_dict()
    if pl.spec_a or pl.spec_b:
        info.confidence["product_list"] = pl.source

    return info


def read_document_texts(upload_dir: Path) -> dict[str, str]:
    texts: dict[str, str] = {}
    for path in upload_dir.rglob("*.docx"):
        if path.is_file():
            try:
                rel = str(path.relative_to(upload_dir))
            except ValueError:
                rel = path.name
            texts[rel] = read_docx_full_text(path)
    return texts
