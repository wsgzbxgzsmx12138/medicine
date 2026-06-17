from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document

from core.extractor import read_docx_full_text
from core.scanner import FileInfo, files_index
from core.utils import glob_files, load_json, normalize_text


@dataclass
class CompletenessIssue:
    rule_id: str
    severity: str
    message: str
    suggestion: str
    regulation_ref: str = ""
    notify: str = "注册事务负责人"
    category: str = "申报资料"


@dataclass
class ConsistencyIssue:
    field: str
    label: str
    severity: str
    values: dict[str, str]
    suggestion: str


@dataclass
class StructureIssue:
    doc_name: str
    missing_sections: list[str]
    severity: str = "warning"


def _match_any(name: str, pattern: str) -> bool:
    norm = normalize_text(name)
    for part in pattern.split("|"):
        if normalize_text(part) in norm:
            return True
    return False


def _default_notify(rules: dict[str, Any]) -> str:
    return rules.get("regulation", {}).get("default_notify", "注册事务负责人")


def _append_missing(
    issues: list[CompletenessIssue],
    *,
    rule_id: str,
    severity: str,
    message: str,
    suggestion: str,
    regulation_ref: str,
    notify: str,
    category: str,
) -> None:
    issues.append(
        CompletenessIssue(
            rule_id=rule_id,
            severity=severity,
            message=message,
            suggestion=suggestion,
            regulation_ref=regulation_ref,
            notify=notify,
            category=category,
        )
    )


def check_completeness(file_list: list[FileInfo], rules: dict[str, Any] | None = None) -> list[CompletenessIssue]:
    """
    对照 CMDE 2021年第121号公告及附件4《体外诊断试剂注册申报资料要求及说明》
    检查上传文件夹中申报资料是否齐全。
    """
    rules = rules or load_json("nmpa_rules.json")
    index = files_index(file_list)
    notify = _default_notify(rules)
    issues: list[CompletenessIssue] = []

    # 1) CMDE 公告附件（图2 七份法规文件；首次注册仅强制 A4/A7，其余按配置）
    for att in rules.get("cmde_attachments", []):
        if not att.get("required_for_initial", False):
            continue
        if _match_any(index, att["match"]):
            continue
        _append_missing(
            issues,
            rule_id=att["id"],
            severity=att.get("severity", "warning"),
            message=att.get("msg", f"缺少 {att.get('title', '')}"),
            suggestion=att.get("suggestion", f"请补充与「{att.get('title', '')}」对应的文件"),
            regulation_ref=att.get("regulation_ref", f"CMDE 公告附件{att.get('attachment_no', '')}"),
            notify=notify,
            category="法规参考文件",
        )

    # 2) 附件4 各章必交申报资料（产品技术要求、临床评价、注册检验等）
    for rule in rules.get("submission_required", []):
        if _match_any(index, rule["match"]):
            continue
        _append_missing(
            issues,
            rule_id=rule["id"],
            severity=rule.get("severity", "critical"),
            message=rule.get("msg", f"缺少 {rule.get('title', '')}"),
            suggestion=rule.get(
                "suggestion",
                f"请补充包含「{rule['match'].split('|')[0]}」关键词的申报文件",
            ),
            regulation_ref=rule.get("regulation_ref", rules.get("regulation", {}).get("checklist_basis", "")),
            notify=notify,
            category="申报资料",
        )

    # 3) 第1章 CH1 监管信息
    for rule in rules.get("ch1_required", []):
        if not rule.get("required", True):
            continue
        if _match_any(index, rule["match"]):
            continue
        msg = f"缺少 {rule['code']} {rule.get('title', '')}"
        if rule.get("conditional"):
            msg += f"（{rule['conditional']}）"
        _append_missing(
            issues,
            rule_id=rule["code"],
            severity=rule.get("severity", "critical"),
            message=msg,
            suggestion=f"请提交 {rule['code']} {rule.get('title', '')} 对应资料文件",
            regulation_ref=rule.get("regulation_ref", "附件4 第1章"),
            notify=notify,
            category="CH1 监管信息",
        )

    # 按严重等级排序：critical > warning > info
    severity_order = {"critical": 0, "warning": 1, "info": 2}
    issues.sort(key=lambda i: (severity_order.get(i.severity, 9), i.rule_id))
    return issues


def extract_table_by_label(doc_path: Path, row_label: str) -> str | None:
    doc = Document(str(doc_path))
    label_norm = normalize_text(row_label)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if not cells:
                continue
            if label_norm in normalize_text(cells[0]) or any(label_norm in normalize_text(c) for c in cells[:2]):
                for cell in cells[1:]:
                    if cell.strip() and normalize_text(cell) != label_norm:
                        return cell.strip()
    return None


def extract_first_heading(doc_path: Path) -> str | None:
    doc = Document(str(doc_path))
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and len(t) > 4 and "产品列表" not in t and "监管信息" not in t:
            return t
    return None


def extract_regex_from_doc(text: str, pattern: str) -> str | None:
    m = re.search(pattern, text, re.S)
    if not m:
        return None
    return m.group(1).strip()


def extract_section_from_text(text: str, section: str) -> str | None:
    pattern = rf"{re.escape(section)}\s*\n?\s*(.+?)(?=\n【|\Z)"
    m = re.search(pattern, text, re.S)
    return m.group(1).strip()[:500] if m else None


def collect_field_values(upload_dir: Path, field_cfg: dict[str, Any]) -> dict[str, str]:
    values: dict[str, str] = {}
    for src in field_cfg.get("sources", []):
        glob_pat = src["doc_glob"]
        paths = glob_files(upload_dir, glob_pat)
        if not paths:
            continue
        path = paths[0]
        method = src.get("extract")
        val: str | None = None

        if method == "regex_section":
            text = read_docx_full_text(path)
            val = extract_section_from_text(text, src["section"])
        elif method == "table_label":
            val = extract_table_by_label(path, src["label"])
        elif method == "first_heading":
            val = extract_first_heading(path)
        elif method == "regex":
            text = read_docx_full_text(path)
            val = extract_regex_from_doc(text, src["pattern"])
        elif method == "regex_line":
            text = read_docx_full_text(path)
            m = re.search(src["pattern"], text)
            val = m.group(1).strip() if m else None

        if val:
            values[path.name] = val
    return values


def check_consistency(upload_dir: Path, mapping: dict[str, Any] | None = None) -> list[ConsistencyIssue]:
    mapping = mapping or load_json("field_mapping.json")
    issues: list[ConsistencyIssue] = []

    for field_cfg in mapping.get("consistency_fields", []):
        values = collect_field_values(upload_dir, field_cfg)
        if len(values) < 2:
            continue
        normalized = {k: normalize_text(v) for k, v in values.items()}
        unique = set(normalized.values())
        if len(unique) > 1:
            issues.append(
                ConsistencyIssue(
                    field=field_cfg["field"],
                    label=field_cfg.get("label", field_cfg["field"]),
                    severity="warning",
                    values=values,
                    suggestion=f"以下文件中的「{field_cfg.get('label')}」不一致，请核对并统一",
                )
            )
    return issues


def check_manual_structure(upload_dir: Path, rules: dict[str, Any] | None = None) -> list[StructureIssue]:
    rules = rules or load_json("nmpa_rules.json")
    required = rules.get("manual_required_sections", [])
    issues: list[StructureIssue] = []
    manual_files = glob_files(upload_dir, "*说明书*.docx")
    for path in manual_files:
        text = read_docx_full_text(path)
        missing = [s for s in required if s not in text]
        if missing:
            issues.append(StructureIssue(doc_name=path.name, missing_sections=missing))
    return issues
